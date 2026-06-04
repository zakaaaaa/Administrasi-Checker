"""
ReferenceValidator — validator Daftar Pustaka sesuai blueprint v0.3 §4.5.

Yang dicek:
A. Format strict Harvard (pengecekan dasar — bukan parser Harvard penuh):
   - Ada tahun di entry
   - Tidak terlalu pendek (entry minimal punya nama + tahun + judul)

B. LARANGAN STRICT (Daftar Pustaka):
   - Tidak boleh "et al." / "dkk." di Daftar Pustaka — wajib tulis semua nama
   - Di sitasi in-text "et al." / "dkk." DIPERBOLEHKAN.

C. Pengecekan keseimbangan sitasi (anti-referensi-bodong):
   1. In-text → DP: tiap sitasi (Author, Year) di body teks WAJIB ada di DP
   2. DP → In-text: tiap entry DP WAJIB pernah dirujuk di body teks
   - Jika tidak cocok persis tapi ada penulis DP + tahun sama dengan ejaan mirip
     (jarak Levenshtein kecil), beri petunjuk "diduga typo".

D. Urutan alfabetis: entry DP berdasarkan nama belakang author pertama

E. Rekomendasi kualitas: jumlah referensi mutakhir (< 10 tahun)
   - Saran (warning), bukan fail

Catatan desain:
- Format Harvard strict yang penuh (penulis dengan inisial, italic untuk judul,
  format jurnal vs buku) sangat rapuh untuk regex. Kita HANYA validasi yang
  jelas: ada tahun, tidak ada et al./dkk. di DP, balance check.
- Pencocokan sitasi pakai nama pertama (last name) dari entry DP. Variasi
  penulisan (typo kecil) → dicoba deteksi dengan jarak edit; typo besar tetap
  mismatch fail.
- Bibliografi dari plugin (Zotero/Mendeley) sering disisipkan sebagai blok
  w:sdt (kontrol konten) di antara heading «DAFTAR PUSTAKA» dan «LAMPIRAN».
  Teks itu tidak selalu muncul di python-docx Document.paragraphs; karena itu
  ekstraksi DP juga menjelajahi sibling OOXML (termasuk w:sdt / w:sdtContent).
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from typing import Optional

from docx.oxml.ns import qn
from lxml import etree

from app.services.docx_parser import DocxParser, NSMAP
from app.services.schema_rules import SchemaRules


# ============================================================================
# Pattern regex
# ============================================================================

# Pola sitasi in-text (format Harvard benar):
#   (Author, 2020)
#   (Author dan Author, 2020)
#   (Author dkk., 2020)        ← et al. di sitasi in-text diperbolehkan
#   (Author et al., 2020)      ← sama
# Year = 4 digit, optional suffix huruf kecil (2020a, 2020b)
_INTEXT_CITATION_RE = re.compile(
    r"\(([^()]+?),\s*(\d{4}[a-z]?)\)",
    re.UNICODE,
)

# Pola sitasi in-text TANPA koma (format salah tapi masih harus dideteksi):
#   (Author 2020)  ← seharusnya (Author, 2020)
# Hanya tangkap jika author dimulai huruf kapital (filter "Gambar", "Tabel", dst di bawah)
_INTEXT_CITATION_NO_COMMA_RE = re.compile(
    r"\(([A-Z][a-zA-ZÀ-ɏ\-]+"
    r"(?:\s+(?:dan|and|&)\s+[A-Z][a-zA-ZÀ-ɏ\-]+)?"
    r"(?:\s+(?:dkk\.?|et\s+al\.?))?)"
    r"\s+(\d{4}[a-z]?)\)",
    re.UNICODE,
)

# Kata pembuka yang bukan nama penulis — filter false positive (Gambar 2020 dst.)
_NO_COMMA_BLOCKLIST = frozenset([
    "GAMBAR", "TABEL", "LAMPIRAN", "FIGURE", "TABLE", "APPENDIX",
    "BAB", "CHAPTER", "HALAMAN", "PAGE", "GRAFIK", "DIAGRAM",
    "BAGAN", "PETA", "FOTO", "LIHAT", "SEE",
])

# Tahun di awal/dalam entry DP (Indonesia: setelah nama+titik)
# Format yang umum di sample real:
#   "Astuti. 2012. Judul..."           ← year setelah . (Indonesia style)
#   "Astuti (2012) Judul..."           ← Harvard
#   "Astuti, 2012, Judul..."           ← varian
_DP_YEAR_RE = re.compile(r"\b(\d{4}[a-z]?)\b")

# Larangan strict — pakai batas kata agar tidak salah positif (mis. "metal")
_ET_AL_RE = re.compile(r"\bet\s+al\.?\b", re.IGNORECASE)
_DKK_RE = re.compile(r"\bdkk\.?\b", re.IGNORECASE)

# Pattern nama pertama dari entry DP. Strategi: kata-kata di awal sampai
# titik atau koma atau "(" pertama.
# Contoh: "Astuti. 2012. ..." → "Astuti"
#         "Smith, J. (2020) ..." → "Smith"
#         "Nyman, Rimma. 2017. ..." → "Nyman"
#         "BPS. 2018. ..." → "BPS"
#         "UU RI No. 20 Tahun 2003" → "UU RI No"  (special, akan kita filter)
_DP_AUTHOR_RE = re.compile(r"^\s*([^.,(\d]+?)(?:[.,(]|\s+\d{4})", re.UNICODE)


def _levenshtein(a: str, b: str) -> int:
    """Jarak edit antara dua string (case-sensitive; caller normalisasi)."""
    la, lb = len(a), len(b)
    if la == 0:
        return lb
    if lb == 0:
        return la
    row = list(range(lb + 1))
    for i in range(1, la + 1):
        prev = row[0]
        row[0] = i
        for j in range(1, lb + 1):
            cur = row[j]
            cost = 0 if a[i - 1] == b[j - 1] else 1
            row[j] = min(row[j] + 1, row[j - 1] + 1, prev + cost)
            prev = cur
    return row[-1]


def _dp_has_et_al(text: str) -> bool:
    return bool(_ET_AL_RE.search(text))


def _dp_has_dkk(text: str) -> bool:
    return bool(_DKK_RE.search(text))


# Data class
# ============================================================================


@dataclass
class ReferenceEntry:
    """Satu entry di Daftar Pustaka."""
    paragraph_index: int
    raw_text: str
    author_first: Optional[str] = None    # last name penulis pertama (untuk balance)
    year: Optional[int] = None
    has_etal_violation: bool = False
    has_dkk_violation: bool = False
    has_year: bool = True
    is_too_short: bool = False
    is_recent: bool = False               # year >= cutoff (cutoff = current_year - 10)


@dataclass
class InTextCitation:
    """Satu sitasi di body teks."""
    paragraph_index: int
    raw_text: str               # mis. "(Astuti, 2012)"
    author: str                 # mis. "Astuti" (penulis pertama setelah normalisasi)
    year: int                   # 2012
    author_raw: str = ""        # substring penulis persis dari dokumen (sebelum koma+tahun)
    has_format_error: bool = False  # True jika tidak pakai koma Harvard: "(Astuti 2012)"


@dataclass
class FormatIssue:
    entry_index: int            # >=0: index entry DP; -1 = pelanggaran sitasi in-text
    text_preview: str
    severity: str               # 'fail' | 'warning'
    issue: str
    paragraph_index: Optional[int] = None  # lokasi (DP atau in-text)


@dataclass
class BalanceFinding:
    direction: str              # 'in_text_not_in_references' | 'in_references_not_in_text'
    citation_or_entry: str
    detail: str
    paragraph_index: Optional[int] = None


@dataclass
class CheckMessage:
    level: str
    text: str


@dataclass
class ReferenceValidationResult:
    status: str                 # 'pass' | 'warning' | 'fail'
    total_entries: int = 0
    valid_entries: int = 0
    entries: list[ReferenceEntry] = field(default_factory=list)
    in_text_citations: list[InTextCitation] = field(default_factory=list)
    format_issues: list[FormatIssue] = field(default_factory=list)
    alphabetical_status: str = "pass"  # 'pass' | 'fail'
    out_of_order_pairs: list[tuple[str, str]] = field(default_factory=list)
    balance_findings: list[BalanceFinding] = field(default_factory=list)
    recent_count: int = 0
    recent_threshold_years: int = 10
    minimum_recommended: int = 8
    quality_message: str = ""
    messages: list[CheckMessage] = field(default_factory=list)
    # Lokasi heading (None = tidak ketemu). Untuk pesan error "judul ada tapi kosong".
    dp_heading_paragraph_index: Optional[int] = None
    lampiran_heading_paragraph_index: Optional[int] = None

    def to_dict(self) -> dict:
        return {
            "status": self.status,
            "total_entries": self.total_entries,
            "valid_entries": self.valid_entries,
            "entries": [
                {
                    "paragraph_index": e.paragraph_index,
                    "raw_text": e.raw_text,
                    "author_first": e.author_first,
                    "year": e.year,
                    "has_etal_violation": e.has_etal_violation,
                    "has_dkk_violation": e.has_dkk_violation,
                    "has_year": e.has_year,
                    "is_too_short": e.is_too_short,
                    "is_recent": e.is_recent,
                }
                for e in self.entries
            ],
            "in_text_citations": [
                {
                    "paragraph_index": c.paragraph_index,
                    "raw_text": c.raw_text,
                    "author": c.author,
                    "author_raw": c.author_raw,
                    "year": c.year,
                    "has_format_error": c.has_format_error,
                }
                for c in self.in_text_citations
            ],
            "format_issues": [
                {
                    "entry_index": fi.entry_index,
                    "text_preview": fi.text_preview,
                    "severity": fi.severity,
                    "issue": fi.issue,
                    "paragraph_index": fi.paragraph_index,
                }
                for fi in self.format_issues
            ],
            "alphabetical_order": {
                "status": self.alphabetical_status,
                "out_of_order": [
                    {"current": p[0], "should_be_after": p[1]}
                    for p in self.out_of_order_pairs
                ],
            },
            "citation_balance": {
                "in_text_not_in_references": [
                    {
                        "citation": f.citation_or_entry,
                        "detail": f.detail,
                        "paragraph_index": f.paragraph_index,
                    }
                    for f in self.balance_findings
                    if f.direction == "in_text_not_in_references"
                ],
                "in_references_not_in_text": [
                    {
                        "reference": f.citation_or_entry,
                        "detail": f.detail,
                        "paragraph_index": f.paragraph_index,
                    }
                    for f in self.balance_findings
                    if f.direction == "in_references_not_in_text"
                ],
            },
            "quality_recommendation": {
                "recent_references_count": self.recent_count,
                "recent_threshold_years": self.recent_threshold_years,
                "minimum_recommended": self.minimum_recommended,
                "message": self.quality_message,
            },
            "section_detection": {
                "daftar_pustaka_heading_paragraph_index": self.dp_heading_paragraph_index,
                "lampiran_heading_paragraph_index": self.lampiran_heading_paragraph_index,
            },
            "messages": [{"level": m.level, "text": m.text} for m in self.messages],
        }


# ============================================================================
# ReferenceValidator
# ============================================================================


class ReferenceValidator:
    """
    Validator Daftar Pustaka.

    Usage:
        parser = DocxParser('proposal.docx')
        schema = get_pkm_kc_proposal_rules()
        result = ReferenceValidator(parser, schema, current_year=2026).check()
    """

    # Limit jumlah karakter preview untuk display
    PREVIEW_CHARS = 100

    # Minimal panjang entry (deteksi entry kosong / placeholder)
    MIN_ENTRY_LENGTH = 30

    def __init__(
        self,
        parser: DocxParser,
        schema: SchemaRules,
        current_year: int = 2026,
        recent_threshold_years: int = 10,
        minimum_recommended_recent: int = 8,
    ):
        self.parser = parser
        self.schema = schema
        self.current_year = current_year
        self.recent_threshold_years = recent_threshold_years
        self.minimum_recommended_recent = minimum_recommended_recent

    def check(self) -> ReferenceValidationResult:
        result = ReferenceValidationResult(
            status="pass",
            recent_threshold_years=self.recent_threshold_years,
            minimum_recommended=self.minimum_recommended_recent,
        )

        # Step 1: ekstrak entries dari Daftar Pustaka
        entries, dp_idx, lamp_idx = self._extract_dp_entries()
        result.entries = entries
        result.dp_heading_paragraph_index = dp_idx
        result.lampiran_heading_paragraph_index = lamp_idx
        result.total_entries = len(result.entries)

        # Step 2: validasi format per entry (et al./dkk., year, length)
        result.format_issues = self._validate_format(result.entries)
        result.valid_entries = len([
            e for e in result.entries
            if not (e.has_etal_violation or e.has_dkk_violation
                    or not e.has_year or e.is_too_short)
        ])

        # Step 3: ekstrak sitasi in-text
        # Catatan: et al./dkk. di sitasi in-text DIPERBOLEHKAN (Harvard standar
        # mengizinkan singkatan untuk 3+ penulis). Yang dilarang hanya di DP.
        result.in_text_citations = self._extract_in_text_citations()

        # Step 3b: flag sitasi in-text yang tidak sesuai format Harvard (tanpa koma)
        for c in result.in_text_citations:
            if c.has_format_error:
                result.format_issues.append(
                    FormatIssue(
                        entry_index=-1,
                        text_preview=c.raw_text,
                        severity="fail",
                        issue=(
                            f"Format sitasi in-text tidak sesuai Harvard: '{c.raw_text}' "
                            f"— nama penulis dan tahun harus dipisah koma. "
                            f"Seharusnya '({c.author}, {c.year})'."
                        ),
                        paragraph_index=c.paragraph_index,
                    )
                )

        # Step 4: balance check (in-text ↔ DP)
        result.balance_findings = self._balance_check(
            result.entries, result.in_text_citations
        )

        # Step 5: alphabetical order check
        result.alphabetical_status, result.out_of_order_pairs = (
            self._check_alphabetical_order(result.entries)
        )

        # Step 6: quality recommendation (recency)
        result.recent_count = sum(1 for e in result.entries if e.is_recent)
        result.quality_message = self._build_quality_message(result.recent_count)

        # Step 7: aggregate status
        self._finalize_status(result)

        return result

    # ------------------------------------------------------------------------
    # Step 1: ekstrak entries
    # ------------------------------------------------------------------------

    def _extract_dp_entries(
        self,
    ) -> tuple[list[ReferenceEntry], Optional[int], Optional[int]]:
        """
        Locate section "DAFTAR PUSTAKA" (heading_only), ambil paragraf
        sampai section selanjutnya (LAMPIRAN). Entry = paragraf non-empty.

        Return (entries, indeks paragraf judul DP atau None, indeks judul LAMPIRAN atau None).
        """
        boundaries = self.parser.find_section_boundaries(
            ["DAFTAR PUSTAKA", "LAMPIRAN"], headings_only=True
        )
        dp_start = boundaries["DAFTAR PUSTAKA"]
        dp_end = boundaries["LAMPIRAN"]

        if dp_start is None:
            return [], None, dp_end

        # Kalau LAMPIRAN tidak ada, ambil sampai akhir
        if dp_end is None:
            dp_end = len(self.parser.paragraphs)

        blocks = self._walk_dp_body_siblings_ooxml(dp_start, dp_end)
        entries: list[ReferenceEntry] = []
        if blocks:
            for para_idx, text_normalized in blocks:
                if not self._is_meaningful_entry(text_normalized):
                    continue
                entry = self._parse_entry(para_idx, text_normalized)
                entries.append(entry)
        else:
            for i in range(dp_start + 1, dp_end):
                para = self.parser.paragraphs[i]
                text = para.text.strip()
                if not text:
                    continue
                if para.is_heading:
                    continue
                text_normalized = re.sub(r"\s+", " ", text)
                if not self._is_meaningful_entry(text_normalized):
                    continue
                entry = self._parse_entry(i, text_normalized)
                entries.append(entry)
        entries = self._merge_continuation_entries(entries)
        return entries, dp_start, dp_end

    def _walk_dp_body_siblings_ooxml(
        self, dp_start: int, dp_end: int
    ) -> list[tuple[int, str]]:
        """
        Kumpulkan teks entry DP dari sibling OOXML antara paragraf judul DP
        dan paragraf judul LAMPIRAN.

        Menangkap isi w:sdt (bibliografi Zotero/Mendeley/Word) yang tidak
        termasuk di Document.paragraphs.
        """
        doc = self.parser.doc
        paras = doc.paragraphs
        if dp_start < 0 or dp_start >= len(paras):
            return []
        try:
            start_el = paras[dp_start]._element
            if dp_end is None or dp_end >= len(paras):
                end_el: Optional[etree._Element] = None
            else:
                end_el = paras[dp_end]._element
        except (AttributeError, IndexError, TypeError):
            return []

        out: list[tuple[int, str]] = []
        default_idx = dp_start + 1

        def norm(s: str) -> str:
            return re.sub(r"\s+", " ", (s or "").strip())

        def text_from_p(p_el: etree._Element) -> str:
            return "".join(
                t.text or "" for t in p_el.findall(".//w:t", namespaces=NSMAP)
            )

        cur: Optional[etree._Element] = start_el.getnext()
        try:
            while cur is not None:
                if end_el is not None and cur is end_el:
                    break
                tag = etree.QName(cur).localname
                if tag == "sectPr":
                    break
                if tag == "p":
                    text = norm(text_from_p(cur))
                    if text:
                        idx = default_idx
                        skip_para = False
                        for pi, p in enumerate(paras):
                            if p._element is cur:
                                idx = pi
                                if (
                                    pi < len(self.parser.paragraphs)
                                    and self.parser.paragraphs[pi].is_heading
                                ):
                                    skip_para = True
                                break
                        if not skip_para:
                            out.append((idx, text))
                elif tag == "sdt":
                    content = cur.find(qn("w:sdtContent"))
                    if content is not None:
                        for inner in content.findall(qn("w:p")):
                            text = norm(text_from_p(inner))
                            if text:
                                out.append((default_idx, text))
                cur = cur.getnext()
        except (AttributeError, TypeError):
            return []
        return out

    def _parse_entry(self, para_idx: int, text: str) -> ReferenceEntry:
        entry = ReferenceEntry(
            paragraph_index=para_idx,
            raw_text=text,
            is_too_short=len(text) < self.MIN_ENTRY_LENGTH,
        )

        # Cek et al./dkk. (batas kata)
        if _dp_has_et_al(text):
            entry.has_etal_violation = True
        if _dp_has_dkk(text):
            entry.has_dkk_violation = True

        # Cari tahun
        m = _DP_YEAR_RE.search(text)
        if m:
            year_str = m.group(1)
            # Buang suffix huruf untuk konversi int
            year_int_str = re.sub(r"[a-z]", "", year_str)
            try:
                entry.year = int(year_int_str)
                # Validasi tahun masuk akal (1900..current_year+1)
                if not (1900 <= entry.year <= self.current_year + 1):
                    entry.year = None
            except ValueError:
                entry.year = None
        if entry.year is None:
            entry.has_year = False

        # Recency
        if entry.year is not None:
            entry.is_recent = (
                self.current_year - entry.year < self.recent_threshold_years
            )

        # Author first (untuk balance check)
        m_author = _DP_AUTHOR_RE.match(text)
        if m_author:
            author = m_author.group(1).strip()
            # Bersihkan dari titik trailing
            author = author.rstrip(". ").strip()
            # Skip kalau terlalu generik (mis. "UU RI No")
            if author and len(author) >= 2:
                if "," in author:
                    # Format "Last, First" → ambil last name saja
                    entry.author_first = author.split(",")[0].strip()
                else:
                    # Nama institusi / single-name — simpan utuh
                    entry.author_first = author

        return entry

    @staticmethod
    def _is_meaningful_entry(text: str) -> bool:
        """Return False untuk teks yang tidak mengandung huruf/angka (mis. satu titik artefak formatting)."""
        return bool(re.search(r"[A-Za-z0-9À-ɏ]", text))

    @staticmethod
    def _looks_like_new_entry(text: str) -> bool:
        """
        Return True jika teks terlihat seperti awal entry baru (bukan lanjutan).

        Continuation paragraph biasanya:
        - Dimulai dengan nama jurnal / info volume:halaman (≥5 kata sebelum delimiter pertama)
        - Dimulai dengan URL / DOI / "Retrieved"
        Sedangkan entry baru selalu dimulai nama pengarang (1-3 kata + koma/titik cepat).
        """
        if re.match(r"^\s*(https?://|doi:|Retrieved|Available)", text, re.IGNORECASE):
            return False
        m = re.match(r"^\s*(\w[\w\s]{0,50}?)[,.(]", text)
        if not m:
            return True
        prefix_words = m.group(1).strip().split()
        # ≥5 kata sebelum delimiter pertama → kemungkinan judul jurnal/buku (continuation)
        return len(prefix_words) < 5

    def _merge_continuation_entries(
        self, entries: list[ReferenceEntry]
    ) -> list[ReferenceEntry]:
        """
        Gabungkan paragraf lanjutan (yang terpotong style/wrap) ke entry sebelumnya.

        Paragraf dianggap lanjutan jika:
        1. Tidak punya tahun, DAN
        2. Tidak terlihat seperti awal entry baru (nama pengarang pendek + delimiter)
        """
        if not entries:
            return entries
        merged: list[ReferenceEntry] = [entries[0]]
        for entry in entries[1:]:
            if not entry.has_year and not self._looks_like_new_entry(entry.raw_text):
                prev = merged[-1]
                combined_text = prev.raw_text + " " + entry.raw_text
                merged[-1] = self._parse_entry(prev.paragraph_index, combined_text)
            else:
                merged.append(entry)
        return merged

    # ------------------------------------------------------------------------
    # Step 2: validasi format
    # ------------------------------------------------------------------------

    def _validate_format(
        self, entries: list[ReferenceEntry]
    ) -> list[FormatIssue]:
        issues: list[FormatIssue] = []
        for i, e in enumerate(entries):
            preview = e.raw_text[: self.PREVIEW_CHARS]
            para_idx = e.paragraph_index
            if e.has_etal_violation or e.has_dkk_violation:
                issues.append(
                    FormatIssue(
                        entry_index=i,
                        text_preview=preview,
                        severity="fail",
                        issue=(
                            "Nama penulis di daftar pustaka tidak boleh ditulis "
                            "dkk atau et al, tuliskan semua nama penulis."
                        ),
                        paragraph_index=para_idx,
                    )
                )
            if not e.has_year:
                issues.append(
                    FormatIssue(
                        entry_index=i,
                        text_preview=preview,
                        severity="fail",
                        issue=(
                            "Tahun publikasi tidak ditemukan di entry. "
                            "Setiap referensi wajib mencantumkan tahun."
                        ),
                        paragraph_index=para_idx,
                    )
                )
            if e.is_too_short:
                issues.append(
                    FormatIssue(
                        entry_index=i,
                        text_preview=preview,
                        severity="warning",
                        issue=(
                            f"Entry terlalu pendek ({len(e.raw_text)} char) "
                            f"— kemungkinan tidak lengkap."
                        ),
                        paragraph_index=para_idx,
                    )
                )
        return issues

    # ------------------------------------------------------------------------
    # Step 3: ekstrak sitasi in-text
    # ------------------------------------------------------------------------

    def _extract_in_text_citations(self) -> list[InTextCitation]:
        """
        Scan body teks dari BAB 1 sampai sebelum DAFTAR PUSTAKA.

        Dua pass:
        1. Pattern benar: (Author, Year)
        2. Pattern salah (tanpa koma): (Author Year) → has_format_error=True
           Tetap dimasukkan ke list agar balance check tidak salah "bodong".
        """
        boundaries = self.parser.find_section_boundaries(
            ["BAB 1", "DAFTAR PUSTAKA"], headings_only=True
        )
        body_start = boundaries["BAB 1"]
        body_end = boundaries["DAFTAR PUSTAKA"]

        if body_start is None:
            body_start = 0
        if body_end is None:
            body_end = len(self.parser.paragraphs)

        citations: list[InTextCitation] = []

        for i in range(body_start, body_end):
            text = self.parser.paragraphs[i].text
            if not text.strip():
                continue

            # Pass 1: format benar (dengan koma)
            matched_spans: set[tuple[int, int]] = set()
            for m in _INTEXT_CITATION_RE.finditer(text):
                author_part = m.group(1).strip()
                year_str = m.group(2)
                year_int_str = re.sub(r"[a-z]", "", year_str)
                try:
                    year = int(year_int_str)
                except ValueError:
                    continue
                first_author = self._extract_first_author_from_citation(author_part)
                if not first_author:
                    continue
                matched_spans.add(m.span())
                citations.append(
                    InTextCitation(
                        paragraph_index=i,
                        raw_text=m.group(0),
                        author=first_author,
                        year=year,
                        author_raw=author_part,
                        has_format_error=False,
                    )
                )

            # Pass 2: format salah (tanpa koma) — (Author Year)
            for m in _INTEXT_CITATION_NO_COMMA_RE.finditer(text):
                if m.span() in matched_spans:
                    continue
                author_part = m.group(1).strip()
                # Filter false positive: Gambar, Tabel, BAB, dst.
                first_word = author_part.split()[0].upper()
                if first_word in _NO_COMMA_BLOCKLIST:
                    continue
                year_str = m.group(2)
                year_int_str = re.sub(r"[a-z]", "", year_str)
                try:
                    year = int(year_int_str)
                except ValueError:
                    continue
                first_author = self._extract_first_author_from_citation(author_part)
                if not first_author:
                    continue
                citations.append(
                    InTextCitation(
                        paragraph_index=i,
                        raw_text=m.group(0),
                        author=first_author,
                        year=year,
                        author_raw=author_part,
                        has_format_error=True,
                    )
                )

        return citations

    @staticmethod
    def _extract_first_author_from_citation(author_text: str) -> Optional[str]:
        """
        Ambil author pertama dari teks sitasi. Toleran terhadap:
        - "Smith"
        - "Smith dan Jones"
        - "Smith & Jones"
        - "Smith et al."
        - "Smith dkk."
        """
        # Buang prefix umum
        text = author_text.strip()
        # Split by 'dan', '&', ',', 'and', dst.
        tokens = re.split(r"\s+(?:dan|and|&)\s+|,\s+", text, maxsplit=1)
        if not tokens:
            return None
        first = tokens[0].strip()
        # Buang trailing "et al" / "dkk" (batas kata)
        first = _ET_AL_RE.sub("", first).strip()
        first = _DKK_RE.sub("", first).strip()
        first = first.strip(". ").strip()
        if not first or len(first) < 2:
            return None
        return first

    # ------------------------------------------------------------------------
    # Step 4: balance check
    # ------------------------------------------------------------------------

    def _balance_check(
        self,
        entries: list[ReferenceEntry],
        citations: list[InTextCitation],
    ) -> list[BalanceFinding]:
        """
        Dua arah cek:
        1. In-text → DP: tiap (Author, Year) di teks harus ada di DP
        2. DP → In-text: tiap entry DP harus pernah disitasi

        Pencocokan: berdasarkan (author_first.lower(), year).
        Toleransi: cek substring author (untuk handle "Mustika" vs "Muatika")
        """
        # Build sets
        dp_keys: set[tuple[str, int]] = set()
        dp_authors_lower: dict[str, list[ReferenceEntry]] = {}
        for e in entries:
            if e.author_first and e.year:
                dp_keys.add((e.author_first.lower(), e.year))
                dp_authors_lower.setdefault(e.author_first.lower(), []).append(e)

        intext_keys: set[tuple[str, int]] = set()
        intext_authors_lower: dict[str, list[InTextCitation]] = {}
        for c in citations:
            intext_keys.add((c.author.lower(), c.year))
            intext_authors_lower.setdefault(c.author.lower(), []).append(c)

        findings: list[BalanceFinding] = []

        seen_intext_findings: set[tuple[str, int]] = set()
        # 1. In-text not in DP
        for c in citations:
            key = (c.author.lower(), c.year)
            if key in dp_keys:
                continue
            # Coba match toleran: cek apakah ada entry DP dengan author yang
            # PREFIX-cocok (mis. "Smith" vs "Smithson")
            partial = self._fuzzy_author_match(c.author.lower(), dp_authors_lower.keys())
            if partial:
                dp_years = {e.year for e in dp_authors_lower[partial] if e.year}
                if c.year not in dp_years:
                    if key in seen_intext_findings:
                        continue
                    seen_intext_findings.add(key)
                    findings.append(
                        BalanceFinding(
                            direction="in_text_not_in_references",
                            citation_or_entry=c.raw_text,
                            detail=(
                                f"Sitasi '({c.author}, {c.year})' tidak sesuai "
                                f"di Daftar Pustaka. Author '{partial}' ada tapi "
                                f"tahun {sorted(dp_years)} (bukan {c.year})."
                            ),
                            paragraph_index=c.paragraph_index,
                        )
                    )
                elif c.author.lower() != partial.lower():
                    # Prefix mirip tapi ejaan beda — jangan anggap sudah valid.
                    if key in seen_intext_findings:
                        continue
                    seen_intext_findings.add(key)
                    findings.append(
                        BalanceFinding(
                            direction="in_text_not_in_references",
                            citation_or_entry=c.raw_text,
                            detail=(
                                f"Sitasi \"({c.author}, {c.year})\" tidak cocok persis "
                                f"dengan nama di Daftar Pustaka \"{partial}\". Periksa ejaan penulis."
                            ),
                            paragraph_index=c.paragraph_index,
                        )
                    )
                continue
            else:
                if key in seen_intext_findings:
                    continue
                seen_intext_findings.add(key)
                detail = (
                    f"Sitasi '({c.author}, {c.year})' di teks tapi tidak "
                    f"ada author '{c.author}' di Daftar Pustaka."
                )
                closest = self._closest_dp_author_same_year(c.author, c.year, entries)
                if closest is not None:
                    dp_name, dist = closest
                    th = self._typo_distance_threshold(len(c.author))
                    if 0 < dist <= th:
                        detail = (
                            f"Sitasi '({c.author}, {c.year})' tidak cocok persis "
                            f"dengan Daftar Pustaka. Diduga typo: di DP ada "
                            f"'{dp_name}' ({c.year}) (jarak ejaan {dist})."
                        )
                findings.append(
                    BalanceFinding(
                        direction="in_text_not_in_references",
                        citation_or_entry=c.raw_text,
                        detail=detail,
                        paragraph_index=c.paragraph_index,
                    )
                )

        seen_dp_findings: set[tuple[str, int]] = set()
        # 2. DP not in in-text
        for e in entries:
            if not e.author_first or not e.year:
                continue
            key = (e.author_first.lower(), e.year)
            if key in intext_keys:
                continue
            if key in seen_dp_findings:
                continue
            seen_dp_findings.add(key)
            # Fuzzy author match
            partial = self._fuzzy_author_match(
                e.author_first.lower(), intext_authors_lower.keys()
            )
            if partial:
                # Author di-cite tapi tahun beda — kita TIDAK flag (sudah
                # di-flag oleh sisi in-text)
                continue
            # Referensi di DP tapi tidak disitasi di body → diperbolehkan, skip.

        return findings

    @staticmethod
    def _typo_distance_threshold(author_len: int) -> int:
        if author_len <= 3:
            return 1
        if author_len <= 6:
            return 2
        return 3

    def _closest_dp_author_same_year(
        self,
        cite_author: str,
        year: int,
        entries: list[ReferenceEntry],
    ) -> Optional[tuple[str, int]]:
        """Return (author_first dari DP, jarak Levenshtein) untuk entry tahun sama; None jika tidak ada."""
        a = cite_author.strip().lower()
        if not a:
            return None
        best: Optional[tuple[str, int]] = None
        for e in entries:
            if e.year != year or not e.author_first:
                continue
            d = _levenshtein(a, e.author_first.lower())
            if best is None or d < best[1]:
                best = (e.author_first, d)
        return best

    @staticmethod
    def _fuzzy_author_match(target: str, candidates) -> Optional[str]:
        """
        Cari candidate yang mirip target (exact atau prefix 4+ karakter).
        Return: candidate yang match, atau None.
        """
        if not target:
            return None
        t = target.lower()
        for cand in candidates:
            c = cand.lower()
            if c == t:
                return cand
            if len(t) >= 4 and len(c) >= 4:
                if c.startswith(t[:4]) or t.startswith(c[:4]):
                    return cand
        return None

    # ------------------------------------------------------------------------
    # Step 5: alphabetical order
    # ------------------------------------------------------------------------

    def _check_alphabetical_order(
        self, entries: list[ReferenceEntry]
    ) -> tuple[str, list[tuple[str, str]]]:
        """
        Cek urutan alfabetis berdasarkan author_first.
        Skip entries yang author_first = None (tidak bisa di-sort).

        Return (status, list of (current_name, should_be_after_name)).
        """
        out_of_order: list[tuple[str, str]] = []
        sortable = [e for e in entries if e.author_first]
        for i in range(1, len(sortable)):
            prev = sortable[i - 1].author_first or ""
            curr = sortable[i].author_first or ""
            if curr.lower() < prev.lower():
                out_of_order.append((curr, prev))
        status = "fail" if out_of_order else "pass"
        return status, out_of_order

    # ------------------------------------------------------------------------
    # Step 6: quality recommendation
    # ------------------------------------------------------------------------

    def _build_quality_message(self, recent_count: int) -> str:
        if recent_count >= self.minimum_recommended_recent:
            return (
                f"Ada {recent_count} referensi mutakhir "
                f"(<{self.recent_threshold_years} tahun) — sudah memenuhi "
                f"rekomendasi minimum {self.minimum_recommended_recent}."
            )
        return (
            f"Hanya {recent_count} referensi mutakhir "
            f"(<{self.recent_threshold_years} tahun). Disarankan menambah "
            f"hingga ≥{self.minimum_recommended_recent} untuk scoring "
            f"maksimal. Bukan error mati."
        )

    # ------------------------------------------------------------------------
    # Step 7: finalize
    # ------------------------------------------------------------------------

    def _finalize_status(self, result: ReferenceValidationResult) -> None:
        def loc(paragraph_index: Optional[int]) -> str:
            if paragraph_index is None:
                return ""
            estimator = getattr(self.parser, "estimate_physical_page", None)
            in_page_estimator = getattr(self.parser, "estimate_paragraph_index_in_page", None)
            page = None
            if callable(estimator):
                raw = estimator(paragraph_index)
                if isinstance(raw, int):
                    page = raw
            in_page = None
            if callable(in_page_estimator):
                raw_in = in_page_estimator(paragraph_index)
                if isinstance(raw_in, int):
                    in_page = raw_in
            if page is None:
                return f" (paragraf #{paragraph_index})"
            if in_page is None or in_page <= 0:
                return f" (halaman fisik ~{page}, paragraf #{paragraph_index})"
            return (
                f" (halaman fisik ~{page}, paragraf ke-{in_page} "
                f"(global #{paragraph_index}))"
            )

        has_fail = False
        has_warning = False

        if result.total_entries == 0:
            result.status = "fail"
            if result.dp_heading_paragraph_index is not None:
                hint = loc(result.dp_heading_paragraph_index)
                result.messages.append(
                    CheckMessage(
                        level="fail",
                        text=(
                            'Judul "DAFTAR PUSTAKA" terdeteksi'
                            f"{hint}, tetapi tidak ada satupun paragraf referensi "
                            'di antara judul itu dan heading "LAMPIRAN" (hanya baris '
                            "kosong atau teks yang tidak dianggap sebagai paragraf "
                            "body). Tulis setiap sumber sebagai paragraf biasa tepat "
                            "di bawah judul. Referensi di dalam tabel/gambar/kotak "
                            "teks tidak ikut terbaca oleh pemeriksa ini."
                        ),
                    )
                )
            else:
                result.messages.append(
                    CheckMessage(
                        level="fail",
                        text=(
                            "Daftar Pustaka tidak ditemukan / kosong. Section ini "
                            "wajib ada untuk PKM."
                        ),
                    )
                )
            return

        # Format issues
        dkk_emitted = False
        for fi in result.format_issues:
            if fi.severity == "fail":
                has_fail = True
            else:
                has_warning = True
            # dkk/et al: dedup — cukup satu pesan tanpa prefix/lokasi
            if "dkk atau et al" in fi.issue:
                if not dkk_emitted:
                    result.messages.append(
                        CheckMessage(level=fi.severity, text=fi.issue)
                    )
                    dkk_emitted = True
                continue
            para = fi.paragraph_index
            if para is None and fi.entry_index >= 0 and fi.entry_index < len(
                result.entries
            ):
                para = result.entries[fi.entry_index].paragraph_index
            prefix = (
                "[Sitasi in-text]"
                if fi.entry_index == -1
                else f"[Entry #{fi.entry_index + 1}]"
            )
            result.messages.append(
                CheckMessage(
                    level=fi.severity,
                    text=f"{prefix} {fi.issue}{loc(para)}",
                )
            )

        # Alphabetical
        if result.alphabetical_status == "fail":
            has_fail = True
            result.messages.append(
                CheckMessage(
                    level="fail",
                    text=(
                        f"Urutan alfabetis Daftar Pustaka tidak sesuai. "
                        f"{len(result.out_of_order_pairs)} pasangan entry "
                        f"tidak terurut."
                    ),
                )
            )
            for cur, prev in result.out_of_order_pairs[:5]:
                result.messages.append(
                    CheckMessage(
                        level="fail",
                        text=f"  • '{cur}' muncul setelah '{prev}' (seharusnya sebelum).",
                    )
                )

        # Balance — hanya cek sitasi in-text yang tidak ada di DP
        in_text_missing = [
            f for f in result.balance_findings
            if f.direction == "in_text_not_in_references"
        ]
        if in_text_missing:
            has_fail = True
            result.messages.append(
                CheckMessage(
                    level="fail",
                    text=(
                        f"{len(in_text_missing)} sitasi di teks tidak ditemukan "
                        f"di Daftar Pustaka."
                    ),
                )
            )
            for f in in_text_missing[:5]:
                result.messages.append(
                    CheckMessage(
                        level="fail",
                        text=f"  • {f.detail}",
                    )
                )

        # Quality recommendation
        if result.recent_count < result.minimum_recommended:
            has_warning = True
            result.messages.append(
                CheckMessage(level="warning", text=result.quality_message)
            )
        else:
            result.messages.append(
                CheckMessage(level="pass", text=result.quality_message)
            )

        # Set status
        if has_fail:
            result.status = "fail"
        elif has_warning:
            result.status = "warning"
        else:
            result.status = "pass"
