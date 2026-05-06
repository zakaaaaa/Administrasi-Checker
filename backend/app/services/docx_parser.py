"""
DocxParser — Foundation layer untuk semua checker module.

Tugas utama:
1. Parse file .docx satu kali, expose data ke checker via interface bersih
2. Berikan akses tingkat tinggi (via python-docx) DAN tingkat rendah (raw OOXML via lxml)
3. Lazy-load XML hanya saat dibutuhkan, cache hasil parse
4. Read-only: parser TIDAK memodifikasi dokumen
5. Toleran terhadap dokumen "weird" — simpan warning di self.warnings, return data partial

Dipakai oleh:
- StructureChecker         (paragraphs, headings, tables)
- PhysicalSheetCounter     (perlu PDF konversi terpisah, tapi parser ini extract teks rujukan section)
- FormatChecker            (sections, runs, font properties)
- PageNumberingChecker     (header_xmls, footer_xmls, sections)
- BudgetAuditor            (tables — extract RAB Bab 4 dan Lampiran 2)
- ReferenceValidator       (paragraphs body teks + section "DAFTAR PUSTAKA")
- BiodataOCRVerifier       (paragraphs section biodata + images)
"""

from __future__ import annotations

import zipfile
from dataclasses import dataclass, field
import re
from pathlib import Path
from typing import Iterator, Optional, Union

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from lxml import etree


# ============================================================================
# Konstanta OOXML namespaces
# ============================================================================

# Namespace paling sering dipakai di OOXML
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
PKG_REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"

NSMAP = {"w": W_NS, "r": R_NS}


# ============================================================================
# Data classes — output struktur untuk konsumsi checker
# ============================================================================


@dataclass
class RunInfo:
    """Informasi satu run (potongan teks dengan formatting seragam)."""
    text: str
    font_name: Optional[str] = None
    font_size_pt: Optional[float] = None  # point (bukan half-point)
    bold: Optional[bool] = None
    italic: Optional[bool] = None
    underline: Optional[bool] = None


@dataclass
class ParagraphInfo:
    """Informasi satu paragraf di body dokumen."""
    index: int                                # urutan paragraf di body (0-based)
    text: str                                 # teks lengkap paragraf
    style_name: Optional[str] = None          # nama style (mis. "Heading 1", "Normal")
    alignment: Optional[str] = None           # 'left', 'center', 'right', 'justify', None
    line_spacing: Optional[float] = None      # spasi baris (1.0, 1.15, 1.5, 2.0, dst.)
    runs: list[RunInfo] = field(default_factory=list)
    is_heading: bool = False                  # heuristik: True jika style 'Heading N' atau pola UPPERCASE
    heading_level: Optional[int] = None       # 1, 2, 3, ... jika is_heading


@dataclass
class TableCellInfo:
    """Informasi satu cell tabel."""
    row: int
    col: int
    text: str
    paragraphs: list[str] = field(default_factory=list)


@dataclass
class TableInfo:
    """Informasi satu tabel."""
    index: int                                # urutan tabel di body (0-based)
    rows: int
    cols: int
    cells: list[TableCellInfo] = field(default_factory=list)
    # Heuristik kasar: tabel ini "kira-kira" tabel apa (RAB? jadwal? biodata?)
    # Diisi nanti oleh checker, parser hanya menyediakan strukturnya
    header_texts: list[str] = field(default_factory=list)  # teks baris pertama


@dataclass
class SectionInfo:
    """
    Informasi satu section dokumen.
    Section di Word = blok dokumen dengan setting halaman/header/footer sendiri.
    """
    index: int
    page_width_dxa: Optional[int] = None      # DXA (1/20 point); A4 = 11906
    page_height_dxa: Optional[int] = None     # A4 = 16838
    margin_top_dxa: Optional[int] = None
    margin_bottom_dxa: Optional[int] = None
    margin_left_dxa: Optional[int] = None
    margin_right_dxa: Optional[int] = None
    page_num_format: Optional[str] = None     # 'decimal', 'lowerRoman', 'upperRoman', dll
    page_num_start: Optional[int] = None      # angka mulai (jika di-restart)
    # Reference ke header/footer (rId yang dipakai untuk lookup di header_xmls/footer_xmls)
    header_refs: dict[str, str] = field(default_factory=dict)  # {'default': 'rId4', 'first': 'rId5', 'even': 'rId6'}
    footer_refs: dict[str, str] = field(default_factory=dict)


@dataclass
class ImageInfo:
    """Informasi gambar di word/media/."""
    filename: str                             # mis. 'image1.png'
    content_type: str                         # mis. 'image/png'
    size_bytes: int


# ============================================================================
# Helper: konversi unit
# ============================================================================


def dxa_to_cm(dxa: Optional[int]) -> Optional[float]:
    """Konversi DXA (1/20 point) ke cm. 1 inch = 1440 DXA = 2.54 cm."""
    if dxa is None:
        return None
    return round(dxa / 1440.0 * 2.54, 3)


def half_points_to_pt(hp: Optional[Union[int, str]]) -> Optional[float]:
    """Konversi half-points (unit OOXML untuk font size) ke point."""
    if hp is None:
        return None
    try:
        return float(hp) / 2.0
    except (ValueError, TypeError):
        return None


# ============================================================================
# DocxParser
# ============================================================================


class DocxParser:
    """
    Parser .docx untuk konsumsi semua checker module.

    Usage:
        parser = DocxParser('/path/to/laporan.docx')
        for para in parser.paragraphs:
            print(para.text, para.style_name)
        for header_part_name, xml_root in parser.header_xmls.items():
            # parse manual untuk page numbering checker
            ...
    """

    def __init__(self, file_path: Union[str, Path]):
        self.file_path = Path(file_path)
        if not self.file_path.exists():
            raise FileNotFoundError(f"File tidak ditemukan: {self.file_path}")
        if not self.file_path.suffix.lower() == ".docx":
            raise ValueError(f"Bukan file .docx: {self.file_path}")

        # Warning yang dikumpulkan selama parsing (parser tidak melempar error)
        self.warnings: list[str] = []

        # Lazy-load: hanya di-init saat property pertama kali diakses
        self._doc: Optional[DocxDocument] = None
        self._paragraphs: Optional[list[ParagraphInfo]] = None
        self._tables: Optional[list[TableInfo]] = None
        self._sections: Optional[list[SectionInfo]] = None
        self._header_xmls: Optional[dict[str, etree._Element]] = None
        self._footer_xmls: Optional[dict[str, etree._Element]] = None
        self._images: Optional[list[ImageInfo]] = None
        self._document_xml: Optional[etree._Element] = None
        self._styles_xml: Optional[etree._Element] = None
        self._numbering_xml: Optional[etree._Element] = None
        self._paragraph_page_estimates: Optional[list[int]] = None
        self._paragraph_index_in_page: Optional[list[int]] = None

    # ------------------------------------------------------------------------
    # Akses dokumen via python-docx (high-level)
    # ------------------------------------------------------------------------

    @property
    def doc(self) -> DocxDocument:
        """Objek python-docx Document (lazy)."""
        if self._doc is None:
            self._doc = Document(str(self.file_path))
        return self._doc

    @property
    def paragraphs(self) -> list[ParagraphInfo]:
        """List semua paragraf di body dokumen."""
        if self._paragraphs is None:
            self._paragraphs = self._extract_paragraphs()
        return self._paragraphs

    @property
    def tables(self) -> list[TableInfo]:
        """List semua tabel di body dokumen."""
        if self._tables is None:
            self._tables = self._extract_tables()
        return self._tables

    @property
    def sections(self) -> list[SectionInfo]:
        """List semua section dokumen (settings halaman + ref ke header/footer)."""
        if self._sections is None:
            self._sections = self._extract_sections()
        return self._sections

    # ------------------------------------------------------------------------
    # Akses raw OOXML (low-level) — untuk PageNumberingChecker
    # ------------------------------------------------------------------------

    @property
    def document_xml(self) -> etree._Element:
        """Raw XML word/document.xml (lazy)."""
        if self._document_xml is None:
            self._document_xml = self._read_xml_part("word/document.xml")
        return self._document_xml

    @property
    def styles_xml(self) -> Optional[etree._Element]:
        """Raw XML word/styles.xml (lazy)."""
        if self._styles_xml is None:
            self._styles_xml = self._read_xml_part("word/styles.xml", required=False)
        return self._styles_xml

    @property
    def numbering_xml(self) -> Optional[etree._Element]:
        """Raw XML word/numbering.xml (lazy, optional — tidak semua docx punya)."""
        if self._numbering_xml is None:
            self._numbering_xml = self._read_xml_part("word/numbering.xml", required=False)
        return self._numbering_xml

    @property
    def header_xmls(self) -> dict[str, etree._Element]:
        """
        Dict {part_name: xml_root} untuk semua word/headerN.xml.
        part_name = nama file di dalam zip, mis. 'word/header1.xml'.
        Dipakai PageNumberingChecker untuk validasi zona awal (romawi pojok kanan bawah).
        """
        if self._header_xmls is None:
            self._header_xmls = self._read_header_footer_xmls(prefix="word/header")
        return self._header_xmls

    @property
    def footer_xmls(self) -> dict[str, etree._Element]:
        """
        Dict {part_name: xml_root} untuk semua word/footerN.xml.
        Dipakai PageNumberingChecker untuk validasi zona awal di footer.
        """
        if self._footer_xmls is None:
            self._footer_xmls = self._read_header_footer_xmls(prefix="word/footer")
        return self._footer_xmls

    @property
    def images(self) -> list[ImageInfo]:
        """List gambar di word/media/. Dipakai BiodataOCRVerifier (deteksi TTD crop)."""
        if self._images is None:
            self._images = self._extract_images()
        return self._images

    # ------------------------------------------------------------------------
    # Helper public
    # ------------------------------------------------------------------------

    def iter_runs(self) -> Iterator[tuple[int, RunInfo]]:
        """
        Iterate semua run di body teks beserta indeks paragrafnya.
        Berguna untuk FormatChecker memeriksa konsistensi font.
        """
        for para in self.paragraphs:
            for run in para.runs:
                yield para.index, run

    def estimate_physical_page(self, paragraph_index: int) -> Optional[int]:
        """
        Estimasi nomor halaman fisik (1-based) untuk suatu paragraf.
        Mengandalkan page break marker di OOXML:
        - w:lastRenderedPageBreak
        - w:br w:type="page"
        - section break non-continuous

        Catatan: ini estimasi terbaik dari metadata DOCX, bukan layout engine penuh.
        """
        if paragraph_index < 0:
            return None
        if self._paragraph_page_estimates is None:
            self._paragraph_page_estimates = self._build_page_estimates()
        if paragraph_index >= len(self._paragraph_page_estimates):
            return None
        return self._paragraph_page_estimates[paragraph_index]

    def estimate_paragraph_index_in_page(self, paragraph_index: int) -> Optional[int]:
        """
        Estimasi urutan paragraf (1-based) di halaman fisik tempat paragraf berada.
        Hanya menghitung paragraf yang non-kosong agar lebih relevan untuk user.
        """
        if paragraph_index < 0:
            return None
        if self._paragraph_index_in_page is None:
            self._paragraph_index_in_page = self._build_paragraph_index_in_page()
        if paragraph_index >= len(self._paragraph_index_in_page):
            return None
        return self._paragraph_index_in_page[paragraph_index]

    def find_section_boundaries(
        self,
        section_names: list[str],
        case_sensitive: bool = False,
        headings_only: bool = False,
    ) -> dict[str, Optional[int]]:
        """
        Cari indeks paragraf awal untuk tiap nama section di list.
        Return dict {section_name: paragraph_index | None}.

        Args:
            section_names: list nama section yang dicari, mis. ['BAB 1', 'DAFTAR PUSTAKA'].
            case_sensitive: jika False (default), pencocokan tidak case-sensitive.
            headings_only: jika True, hanya match paragraf yang terdeteksi sebagai
                heading (is_heading=True). Berguna untuk menghindari false-match
                dari baris di Daftar Isi yang teksnya juga dimulai dengan "BAB 1".
                Default False untuk backward compatibility (StructureChecker mungkin
                justru butuh deteksi entri ToC).

        Contoh:
            # Tanpa headings_only — bisa match baris ToC "BAB 1. ........... 1"
            parser.find_section_boundaries(['BAB 1'])

            # Dengan headings_only — hanya match heading "BAB 1. PENDAHULUAN" asli
            parser.find_section_boundaries(['BAB 1'], headings_only=True)

        Pencocokan: paragraf yang teksnya DIMULAI dengan section_name (setelah strip).
        Yang ditemukan adalah kemunculan PERTAMA yang memenuhi kriteria.
        """
        result: dict[str, Optional[int]] = {name: None for name in section_names}

        # Pass 1: perilaku normal
        for para in self.paragraphs:
            if headings_only and not para.is_heading:
                continue
            text = para.text.strip()
            cmp_text = text if case_sensitive else text.upper()
            for name in section_names:
                if result[name] is not None:
                    continue  # sudah ketemu yang pertama
                cmp_name = name if case_sensitive else name.upper()
                if cmp_text.startswith(cmp_name):
                    result[name] = para.index

        # Pass 2 (fallback): jika headings_only=True dan ada section belum ketemu,
        # coba cari baris non-heading yang "mirip heading nyata" dan bukan baris ToC.
        if headings_only and any(v is None for v in result.values()):
            missing = {k for k, v in result.items() if v is None}
            for para in self.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                if self._looks_like_toc_entry(text):
                    continue
                if not self._looks_like_heading_fallback(text):
                    continue
                cmp_text = text if case_sensitive else text.upper()
                for name in list(missing):
                    cmp_name = name if case_sensitive else name.upper()
                    if cmp_text.startswith(cmp_name):
                        result[name] = para.index
                        missing.remove(name)
                        if not missing:
                            break
                if not missing:
                    break
        return result

    @staticmethod
    def _looks_like_toc_entry(text: str) -> bool:
        """
        Heuristik baris daftar isi/lampiran:
        - ada tab + nomor halaman Romawi/angka di akhir
        - atau ada dot leader + nomor halaman di akhir
        """
        t = text.strip()
        if not t:
            return False
        if re.search(r"\t+\s*([ivxlcdm]+|\d+)\s*$", t, flags=re.IGNORECASE):
            return True
        if re.search(r"[.\s]{4,}\s*(\d+|[ivxlcdm]+)\s*$", t, flags=re.IGNORECASE):
            return True
        return False

    @staticmethod
    def _looks_like_heading_fallback(text: str) -> bool:
        """
        Heuristik heading saat style Heading tidak dipakai:
        - Pola BAB N...
        - Atau mayoritas huruf uppercase (mis. DAFTAR PUSTAKA, LAMPIRAN)
        """
        t = text.strip()
        if not t:
            return False
        if re.match(r"^BAB\s+[IVXLCM0-9]+(?:[.\s]|$)", t, flags=re.IGNORECASE):
            return True
        letters = [c for c in t if c.isalpha()]
        if not letters:
            return False
        upper_ratio = sum(1 for c in letters if c.isupper()) / len(letters)
        return upper_ratio >= 0.8

    def read_raw_part(self, part_name: str) -> Optional[bytes]:
        """
        Baca bytes mentah dari part di zip .docx.
        Berguna kalau ada checker yang butuh akses XML/file lain di paket.
        """
        try:
            with zipfile.ZipFile(self.file_path, "r") as z:
                return z.read(part_name)
        except KeyError:
            return None
        except zipfile.BadZipFile:
            self.warnings.append(f"Bad zip: {self.file_path}")
            return None

    # ------------------------------------------------------------------------
    # Implementasi internal
    # ------------------------------------------------------------------------

    def _extract_paragraphs(self) -> list[ParagraphInfo]:
        result: list[ParagraphInfo] = []
        for idx, para in enumerate(self.doc.paragraphs):
            info = self._paragraph_to_info(idx, para)
            result.append(info)
        return result

    def _build_page_estimates(self) -> list[int]:
        page = 1
        estimates: list[int] = []
        for para in self.doc.paragraphs:
            estimates.append(page)
            # Explicit page break di run.
            page += len(
                para._element.findall(
                    ".//w:br[@w:type='page']",
                    namespaces=NSMAP,
                )
            )
            # Last rendered page break (umum pada dokumen dari Word).
            page += len(
                para._element.findall(
                    ".//w:lastRenderedPageBreak",
                    namespaces=NSMAP,
                )
            )
            # Section break default = next page (kecuali continuous).
            sect_pr = para._element.find(".//w:pPr/w:sectPr", namespaces=NSMAP)
            if sect_pr is not None:
                sect_type = sect_pr.find("w:type", namespaces=NSMAP)
                val = sect_type.get(qn("w:val")) if sect_type is not None else None
                if val is None or str(val).lower() != "continuous":
                    page += 1
        return estimates

    def _build_paragraph_index_in_page(self) -> list[int]:
        if self._paragraph_page_estimates is None:
            self._paragraph_page_estimates = self._build_page_estimates()

        counts_per_page: dict[int, int] = {}
        result: list[int] = []
        for idx, para in enumerate(self.paragraphs):
            page = self._paragraph_page_estimates[idx]
            if not para.text.strip():
                # Paragraf kosong tidak menambah hitungan "ke-n" user-facing.
                result.append(counts_per_page.get(page, 0))
                continue
            counts_per_page[page] = counts_per_page.get(page, 0) + 1
            result.append(counts_per_page[page])
        return result

    def _paragraph_to_info(self, idx: int, para: Paragraph) -> ParagraphInfo:
        # Style name
        style_name: Optional[str] = None
        try:
            if para.style is not None:
                style_name = para.style.name
        except Exception as e:
            self.warnings.append(f"Gagal baca style paragraf {idx}: {e}")

        # Alignment
        alignment_str: Optional[str] = None
        try:
            align = para.paragraph_format.alignment
            if align is not None:
                # WD_ALIGN_PARAGRAPH enum → string. Format: "CENTER (1)" → "center"
                raw = str(align)
                # Ambil kata sebelum spasi atau setelah titik terakhir
                if "." in raw:
                    raw = raw.rsplit(".", 1)[-1]
                if " " in raw:
                    raw = raw.split(" ", 1)[0]
                alignment_str = raw.lower()
        except Exception:
            pass

        # Line spacing
        line_spacing: Optional[float] = None
        try:
            ls = para.paragraph_format.line_spacing
            if ls is not None:
                line_spacing = float(ls)
        except Exception:
            pass

        # Runs
        runs: list[RunInfo] = []
        for run in para.runs:
            run_info = RunInfo(
                text=run.text or "",
                font_name=self._safe_font_name(run),
                font_size_pt=self._safe_font_size(run),
                bold=run.bold,
                italic=run.italic,
                underline=run.underline,
            )
            runs.append(run_info)

        # Heading detection: cek style name dulu, fallback ke heuristik UPPERCASE
        is_heading = False
        heading_level: Optional[int] = None
        if style_name:
            sn_lower = style_name.lower()
            if sn_lower.startswith("heading"):
                is_heading = True
                # Coba ekstrak level dari nama style ("Heading 1" → 1)
                parts = style_name.split()
                if len(parts) > 1 and parts[-1].isdigit():
                    heading_level = int(parts[-1])
                else:
                    heading_level = 1
            elif sn_lower in ("title", "subtitle"):
                is_heading = True
                heading_level = 0  # judul

        return ParagraphInfo(
            index=idx,
            text=para.text or "",
            style_name=style_name,
            alignment=alignment_str,
            line_spacing=line_spacing,
            runs=runs,
            is_heading=is_heading,
            heading_level=heading_level,
        )

    def _safe_font_name(self, run) -> Optional[str]:
        """python-docx kadang return None walau XML punya font. Cek manual ke XML run."""
        try:
            if run.font.name:
                return run.font.name
        except Exception:
            pass
        # Fallback: cek XML rPr/rFonts
        try:
            rpr = run._element.find(qn("w:rPr"))
            if rpr is not None:
                rfonts = rpr.find(qn("w:rFonts"))
                if rfonts is not None:
                    # Coba beberapa atribut umum
                    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
                        val = rfonts.get(qn(f"w:{attr}"))
                        if val:
                            return val
        except Exception:
            pass
        return None

    def _safe_font_size(self, run) -> Optional[float]:
        """Return font size dalam point. python-docx return Pt object atau None."""
        try:
            size = run.font.size
            if size is not None:
                # size adalah objek Emu/Length; .pt mengembalikan float
                return float(size.pt)
        except Exception:
            pass
        # Fallback: parse XML
        try:
            rpr = run._element.find(qn("w:rPr"))
            if rpr is not None:
                sz = rpr.find(qn("w:sz"))
                if sz is not None:
                    val = sz.get(qn("w:val"))
                    if val:
                        return half_points_to_pt(val)
        except Exception:
            pass
        return None

    def _extract_tables(self) -> list[TableInfo]:
        result: list[TableInfo] = []
        for idx, tbl in enumerate(self.doc.tables):
            info = self._table_to_info(idx, tbl)
            result.append(info)
        return result

    def _table_to_info(self, idx: int, tbl: Table) -> TableInfo:
        rows = len(tbl.rows)
        cols = len(tbl.columns) if rows > 0 else 0

        cells: list[TableCellInfo] = []
        header_texts: list[str] = []

        for r_idx, row in enumerate(tbl.rows):
            for c_idx, cell in enumerate(row.cells):
                # Cell.text bisa duplicate jika ada merge — kita ambil apa adanya
                cell_text = cell.text or ""
                cell_paragraphs = [p.text for p in cell.paragraphs]
                cells.append(
                    TableCellInfo(
                        row=r_idx, col=c_idx, text=cell_text, paragraphs=cell_paragraphs
                    )
                )
                if r_idx == 0:
                    header_texts.append(cell_text)

        return TableInfo(
            index=idx, rows=rows, cols=cols, cells=cells, header_texts=header_texts
        )

    def _extract_sections(self) -> list[SectionInfo]:
        """
        Extract section info dengan membaca langsung dari XML <w:sectPr>.

        Pendekatan ini lebih reliable daripada lewat python-docx wrapper:
        wrapper kadang return objek Twips/Emu yang konversi int()-nya tidak
        konsisten antar versi python-docx — nilai bisa terbaca 0 di dokumen
        real meski XML-nya valid. XML attributes pgMar dan pgSz langsung
        memberi DXA sebagai string, jadi konversi tinggal int().
        """
        result: list[SectionInfo] = []
        for idx, sect in enumerate(self.doc.sections):
            sect_pr = sect._sectPr
            info = SectionInfo(index=idx)
            self._populate_section_from_xml(info, sect_pr)
            result.append(info)
        return result

    def _populate_section_from_xml(
        self, info: SectionInfo, sect_pr: Optional[etree._Element]
    ) -> None:
        """Isi SectionInfo dari elemen <w:sectPr> via parsing XML langsung."""
        if sect_pr is None:
            return

        # Page size: <w:pgSz w:w="11906" w:h="16838"/>
        pg_sz = sect_pr.find(qn("w:pgSz"))
        if pg_sz is not None:
            info.page_width_dxa = self._safe_int(pg_sz.get(qn("w:w")))
            info.page_height_dxa = self._safe_int(pg_sz.get(qn("w:h")))

        # Margin: <w:pgMar w:top="..." w:right="..." w:bottom="..." w:left="..."/>
        pg_mar = sect_pr.find(qn("w:pgMar"))
        if pg_mar is not None:
            info.margin_top_dxa = self._safe_int(pg_mar.get(qn("w:top")))
            info.margin_right_dxa = self._safe_int(pg_mar.get(qn("w:right")))
            info.margin_bottom_dxa = self._safe_int(pg_mar.get(qn("w:bottom")))
            info.margin_left_dxa = self._safe_int(pg_mar.get(qn("w:left")))

        # Page number type: <w:pgNumType w:fmt="lowerRoman" w:start="1"/>
        pg_num_type = sect_pr.find(qn("w:pgNumType"))
        if pg_num_type is not None:
            info.page_num_format = pg_num_type.get(qn("w:fmt"))
            start = pg_num_type.get(qn("w:start"))
            if start and start.lstrip("-").isdigit():
                info.page_num_start = int(start)

        # Header references: <w:headerReference w:type="default" r:id="rId4"/>
        for ref in sect_pr.findall(qn("w:headerReference")):
            ref_type = ref.get(qn("w:type")) or "default"
            ref_id = ref.get(qn("r:id"))
            if ref_id:
                info.header_refs[ref_type] = ref_id

        for ref in sect_pr.findall(qn("w:footerReference")):
            ref_type = ref.get(qn("w:type")) or "default"
            ref_id = ref.get(qn("r:id"))
            if ref_id:
                info.footer_refs[ref_type] = ref_id

    @staticmethod
    def _safe_int(val: Optional[str]) -> Optional[int]:
        """Konversi string XML attribute ke int, atau None jika tidak valid."""
        if val is None:
            return None
        try:
            return int(val)
        except (ValueError, TypeError):
            return None

    def _read_xml_part(
        self, part_name: str, required: bool = True
    ) -> Optional[etree._Element]:
        raw = self.read_raw_part(part_name)
        if raw is None:
            if required:
                self.warnings.append(f"Part wajib tidak ditemukan: {part_name}")
            return None
        try:
            return etree.fromstring(raw)
        except etree.XMLSyntaxError as e:
            self.warnings.append(f"Gagal parse XML {part_name}: {e}")
            return None

    def _read_header_footer_xmls(self, prefix: str) -> dict[str, etree._Element]:
        """Baca semua part yang namanya diawali prefix (mis. 'word/header')."""
        result: dict[str, etree._Element] = {}
        try:
            with zipfile.ZipFile(self.file_path, "r") as z:
                for name in z.namelist():
                    if name.startswith(prefix) and name.endswith(".xml"):
                        try:
                            raw = z.read(name)
                            result[name] = etree.fromstring(raw)
                        except (etree.XMLSyntaxError, KeyError) as e:
                            self.warnings.append(f"Gagal parse {name}: {e}")
        except zipfile.BadZipFile:
            self.warnings.append(f"Bad zip saat baca {prefix}*.xml")
        return result

    def _extract_images(self) -> list[ImageInfo]:
        """List file di word/media/."""
        result: list[ImageInfo] = []
        try:
            with zipfile.ZipFile(self.file_path, "r") as z:
                for name in z.namelist():
                    if name.startswith("word/media/"):
                        info = z.getinfo(name)
                        ext = Path(name).suffix.lower().lstrip(".")
                        # Mapping ekstensi → content type sederhana
                        ct_map = {
                            "png": "image/png",
                            "jpg": "image/jpeg",
                            "jpeg": "image/jpeg",
                            "gif": "image/gif",
                            "bmp": "image/bmp",
                            "tiff": "image/tiff",
                            "tif": "image/tiff",
                            "svg": "image/svg+xml",
                            "emf": "image/x-emf",
                            "wmf": "image/x-wmf",
                        }
                        content_type = ct_map.get(ext, f"application/{ext}")
                        result.append(
                            ImageInfo(
                                filename=Path(name).name,
                                content_type=content_type,
                                size_bytes=info.file_size,
                            )
                        )
        except zipfile.BadZipFile:
            self.warnings.append("Bad zip saat extract images")
        return result

    # ------------------------------------------------------------------------
    # Debug / introspection
    # ------------------------------------------------------------------------

    def summary(self) -> dict:
        """Ringkasan cepat untuk debugging / smoke test."""
        return {
            "file": str(self.file_path),
            "paragraph_count": len(self.paragraphs),
            "table_count": len(self.tables),
            "section_count": len(self.sections),
            "header_part_count": len(self.header_xmls),
            "footer_part_count": len(self.footer_xmls),
            "image_count": len(self.images),
            "warnings": list(self.warnings),
        }