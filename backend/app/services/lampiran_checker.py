"""
LampiranChecker — validasi kelengkapan lampiran wajib (PKM-KC & PKM-VGK).

Strategi 3-stage:
    Stage 1. Cek heading "DAFTAR LAMPIRAN" ada di dokumen.
             - Tidak ada → return fail dengan SATU pesan "DAFTAR LAMPIRAN
               tidak ditemukan". Stop, tidak enumerasi lampiran.
    Stage 2. Untuk tiap lampiran wajib, cek body section Lampiran (setelah
             heading "LAMPIRAN" body):
             a) Scan teks paragraf — heading "Lampiran N ..." + kata kunci.
             b) Kalau (a) gagal, fallback OCR (Google Vision via LampiranOcrIndex)
                pada gambar di body Lampiran — sering pakai kalau biodata/surat
                pernyataan di-scan jadi image.
    Stage 3. Untuk lampiran yang ADA di body tapi TIDAK terdaftar di blok
             "Daftar Lampiran" depan → flag mismatch.

Pesan output (tiga jenis):
    (a) "DAFTAR LAMPIRAN tidak ditemukan"
    (b) 'Tidak ditemukan Lampiran N "<label>" pada halaman lampiran'
    (c) 'Kesalahan kelengkapan Daftar Lampiran, tidak ditemukan "<label>"'

Aturan kombinasi:
    - Kalau Daftar Lampiran tidak ada → hanya (a). Tidak enumerasi (b)/(c).
    - Kalau lampiran X missing di body → hanya (b). Skip (c) walau di daftar
      juga missing (akar masalah di body, bukan kelengkapan daftar).
    - Kalau lampiran X ada di body tapi tidak di daftar → (c).
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from typing import Optional

from app.services.docx_parser import DocxParser


# =============================================================================
# Konfigurasi lampiran wajib per skema
# =============================================================================

# Tuple format: (nomor, kata_kunci_judul, label_tampil)
# Kata kunci = daftar substring lower-case yang HARUS muncul di window 200 char
# setelah anchor "Lampiran N" (untuk body text) / di blok daftar.

# PKM-KC (6 lampiran sesuai panduan §C.Sistematika PKM-KC-2026):
#   1. Biodata Ketua/Anggota+Dosen (gabungan, BUKAN dipecah)
#   2. Justifikasi Anggaran Kegiatan
#   3. Susunan Tim Pengusul & Pembagian Tugas
#   4. Surat Pernyataan Ketua Tim Pengusul
#   5. Gambaran Teknologi yang akan Dikembangkan  ← khas KC
#   6. Hasil Uji Periksa Similaritas Proposal
_REQUIRED_LAMPIRAN_KC: list[tuple[int, list[str], str]] = [
    (1, ["biodata", "ketua", "dosen"],                   "Biodata Ketua dan Anggota, serta Dosen Pendamping"),
    (2, ["justifikasi", "anggaran"],                     "Justifikasi Anggaran Kegiatan"),
    (3, ["susunan", "tim", "pengusul", "pembagian"],     "Susunan Tim Pengusul dan Pembagian Tugas"),
    (4, ["pernyataan", "ketua"],                         "Surat Pernyataan Ketua Tim Pengusul"),
    (5, ["gambaran", "teknologi"],                       "Gambaran Teknologi yang akan Dikembangkan"),
    (6, ["uji", "similar"],                              "Hasil Uji Periksa Similaritas Proposal"),
]

_REQUIRED_LAMPIRAN_VGK: list[tuple[int, list[str], str]] = [
    (1, ["biodata", "ketua", "dosen"],                   "Biodata Ketua dan Anggota, serta Dosen Pendamping"),
    (2, ["justifikasi", "anggaran"],                     "Justifikasi Anggaran Kegiatan"),
    (3, ["susunan", "tim", "pengusul", "pembagian"],     "Susunan Tim Pengusul dan Pembagian Tugas"),
    (4, ["pernyataan", "ketua"],                         "Surat Pernyataan Ketua Tim Pengusul"),
    (5, ["uji", "similar"],                              "Hasil Uji Periksa Similaritas Proposal"),
]

# PKM-K: struktur lampiran IDENTIK dengan PKM-VGK (5 lampiran, biodata digabung).
# → tidak perlu list terpisah; factory for_pkm_k me-reuse _REQUIRED_LAMPIRAN_VGK.

# PKM-RE (5 lampiran sesuai panduan §C.Sistematika PKM-RE-2026):
#   1. Biodata Ketua/Anggota+Dosen (gabungan)
#   2. Justifikasi Anggaran Kegiatan
#   3. Susunan Tim Pengusul & Pembagian Tugas
#   4. Surat Pernyataan Ketua Tim Pengusul
#   5. Hasil Uji Periksa Similaritas Proposal
# TIDAK ada "Gambaran Teknologi" (khas KC) atau "Kuisioner" (khas RSH).
_REQUIRED_LAMPIRAN_RE: list[tuple[int, list[str], str]] = [
    (1, ["biodata", "ketua", "dosen"],                   "Biodata Ketua dan Anggota, serta Dosen Pendamping"),
    (2, ["justifikasi", "anggaran"],                     "Justifikasi Anggaran Kegiatan"),
    (3, ["susunan", "tim", "pengusul", "pembagian"],     "Susunan Tim Pengusul dan Pembagian Tugas"),
    (4, ["pernyataan", "ketua"],                         "Surat Pernyataan Ketua Tim Pengusul"),
    (5, ["uji", "similar"],                              "Hasil Uji Periksa Similaritas Proposal"),
]

# PKM-RSH (6 lampiran sesuai panduan §C.Sistematika PKM-RSH-2026):
# Sama dengan RE + tambahan Lampiran 5 "Kuisioner/Pedoman Wawancara yang digunakan".
# Anchored keywords: "kuisi" (cover "Kuisioner"/"Kuesioner") + "wawancara".
_REQUIRED_LAMPIRAN_RSH: list[tuple[int, list[str], str]] = [
    (1, ["biodata", "ketua", "dosen"],                   "Biodata Ketua dan Anggota, serta Dosen Pendamping"),
    (2, ["justifikasi", "anggaran"],                     "Justifikasi Anggaran Kegiatan"),
    (3, ["susunan", "tim", "pengusul", "pembagian"],     "Susunan Tim Pengusul dan Pembagian Tugas"),
    (4, ["pernyataan", "ketua"],                         "Surat Pernyataan Ketua Tim Pengusul"),
    (5, ["kuisi", "wawancara"],                          "Kuisioner/Pedoman Wawancara yang digunakan"),
    (6, ["uji", "similar"],                              "Hasil Uji Periksa Similaritas Proposal"),
]

# PKM-KI (6 lampiran): VGK + "Gambaran Konsep Karya Inovatif" sebelum uji similaritas.
_REQUIRED_LAMPIRAN_KI: list[tuple[int, list[str], str]] = [
    (1, ["biodata", "ketua", "dosen"],                   "Biodata Ketua dan Anggota, serta Dosen Pendamping"),
    (2, ["justifikasi", "anggaran"],                     "Justifikasi Anggaran Kegiatan"),
    (3, ["susunan", "tim", "pengusul", "pembagian"],     "Susunan Tim Pengusul dan Pembagian Tugas"),
    (4, ["pernyataan", "ketua"],                         "Surat Pernyataan Ketua Tim Pengusul"),
    (5, ["konsep", "karya", "inovatif"],                 "Gambaran Konsep Karya Inovatif yang akan Dihasilkan"),
    (6, ["uji", "similar"],                              "Hasil Uji Periksa Similaritas Proposal"),
]

# PKM-PI (8 lampiran): + Surat Pernyataan Kesediaan Mitra, Gambaran IPTEK,
# Denah Lokasi Mitra.
_REQUIRED_LAMPIRAN_PI: list[tuple[int, list[str], str]] = [
    (1, ["biodata", "ketua", "dosen"],                   "Biodata Ketua dan Anggota, serta Dosen Pendamping"),
    (2, ["justifikasi", "anggaran"],                     "Justifikasi Anggaran Kegiatan"),
    (3, ["susunan", "tim", "pengusul", "pembagian"],     "Susunan Tim Pengusul dan Pembagian Tugas"),
    (4, ["pernyataan", "ketua"],                         "Surat Pernyataan Ketua Tim Pengusul"),
    (5, ["kesediaan", "mitra"],                          "Surat Pernyataan Kesediaan Bekerjasama dari Mitra"),
    (6, ["gambaran", "iptek"],                           "Gambaran IPTEK yang akan Diterapkan"),
    (7, ["denah", "lokasi", "mitra"],                    "Denah Detail Lokasi Mitra Program"),
    (8, ["uji", "similar"],                              "Hasil Uji Periksa Similaritas Proposal"),
]

# PKM-PM (7 lampiran): + Surat Pernyataan Kesediaan Mitra, Denah Lokasi Mitra.
_REQUIRED_LAMPIRAN_PM: list[tuple[int, list[str], str]] = [
    (1, ["biodata", "ketua", "dosen"],                   "Biodata Ketua dan Anggota, serta Dosen Pendamping"),
    (2, ["justifikasi", "anggaran"],                     "Justifikasi Anggaran Kegiatan"),
    (3, ["susunan", "tim", "pengusul", "pembagian"],     "Susunan Tim Pengusul dan Pembagian Tugas"),
    (4, ["pernyataan", "ketua"],                         "Surat Pernyataan Ketua Tim Pengusul"),
    (5, ["kesediaan", "mitra"],                          "Surat Pernyataan Kesediaan Bekerja Sama dari Mitra"),
    (6, ["denah", "lokasi", "mitra"],                    "Denah Detail Lokasi Mitra Program"),
    (7, ["uji", "similar"],                              "Hasil Uji Periksa Similaritas Proposal"),
]

# PKM-AI (5 lampiran): naskah artikel ilmiah tanpa Daftar Isi/Daftar Lampiran —
# checker harus skip stage 1 (cross-check Daftar Lampiran), hanya cek body.
_REQUIRED_LAMPIRAN_AI: list[tuple[int, list[str], str]] = [
    (1, ["biodata", "ketua", "dosen"],                   "Biodata Ketua dan Anggota, serta Dosen Pendamping"),
    (2, ["kontribusi", "dosen"],                         "Kontribusi Ketua, Anggota, dan Dosen Pendamping"),
    (3, ["pernyataan", "ketua"],                         "Surat Pernyataan Ketua Tim Penyusun"),
    (4, ["pernyataan", "sumber"],                        "Surat Pernyataan Sumber Tulisan"),
    (5, ["uji", "similar"],                              "Hasil Uji Periksa Similaritas Artikel"),
]

# PKM-GFT (4 lampiran): PKM insentif tanpa pelaksanaan — tidak ada justifikasi
# anggaran. Daftar Lampiran tidak terpisah (gabung di Daftar Isi), sehingga
# checker harus skip stage 1 — gunakan require_daftar=False di factory.
_REQUIRED_LAMPIRAN_GFT: list[tuple[int, list[str], str]] = [
    (1, ["biodata", "ketua", "dosen"],                   "Biodata Ketua dan Anggota, serta Dosen Pendamping"),
    (2, ["susunan", "tim", "pengusul", "pembagian"],     "Susunan Tim Pengusul dan Pembagian Tugas"),
    (3, ["pernyataan", "ketua"],                         "Surat Pernyataan Ketua Tim Pengusul"),
    (4, ["uji", "similar"],                              "Hasil Uji Periksa Similaritas Proposal"),
]


# =============================================================================
# Regex
# =============================================================================

# Heading LAMPIRAN (batas awal section lampiran di body)
_LAMPIRAN_SECTION_RE = re.compile(r"^\s*LAMPIRAN\s*$", re.IGNORECASE)
# Heading "DAFTAR LAMPIRAN" / "DAFTAR LAMPIRAN-LAMPIRAN" di depan dokumen
_DAFTAR_LAMPIRAN_RE = re.compile(r"^\s*daftar\s+lampiran(?:-lampiran)?\s*$", re.IGNORECASE)
# Pattern "Lampiran N" — N angka atau romawi
_LAMPIRAN_N_RE = re.compile(r"\blampiran\s+(\d+|[IVXLC]+)\b", re.IGNORECASE)

# Frasa penanda "Surat Pernyataan Komitmen Tambahan Pendanaan" (istilah panduan)
# sekaligus variasi "surat keterangan pendanaan institusi lain" yang lazim ditulis
# pengusul. Lampiran ini WAJIB hanya bila proposal mencantumkan dana Instansi Lain.
# Cocok bila salah satu frasa muncul utuh di korpus lampiran (teks atau OCR).
_FUNDING_LETTER_PHRASES: list[str] = [
    "komitmen tambahan pendanaan",
    "komitmen pendanaan",
    "tambahan pendanaan dari",
    "keterangan pendanaan",
    "pernyataan pendanaan",
    "pendanaan institusi lain",
    "pendanaan instansi lain",
    "pendanaan dari institusi lain",
    "pendanaan dari instansi lain",
]
_FUNDING_LETTER_LABEL = "Surat Pernyataan Komitmen Tambahan Pendanaan dari Instansi Lain"


# =============================================================================
# Data classes
# =============================================================================


@dataclass
class CheckMessage:
    level: str
    text: str


@dataclass
class LampiranCheckResult:
    status: str
    messages: list[CheckMessage] = field(default_factory=list)

    def to_dict(self) -> dict:
        return {
            "status": self.status,
            "messages": [{"level": m.level, "text": m.text} for m in self.messages],
        }


# =============================================================================
# LampiranChecker
# =============================================================================


class LampiranChecker:
    def __init__(
        self,
        parser: DocxParser,
        required: list[tuple[int, list[str], str]],
        schema_label: str,
        require_daftar: bool = True,
    ):
        self.parser = parser
        self.required = required
        self.schema_label = schema_label
        # Kalau False (PKM-AI), skip Stage 1 (cross-check Daftar Lampiran).
        # Dokumen artikel ilmiah tidak memuat Daftar Isi/Daftar Lampiran —
        # validasi hanya pada body section LAMPIRAN.
        self.require_daftar = require_daftar

    @classmethod
    def for_pkm_kc(cls, parser: DocxParser) -> "LampiranChecker":
        return cls(parser, _REQUIRED_LAMPIRAN_KC, "PKM-KC")

    @classmethod
    def for_pkm_vgk(cls, parser: DocxParser) -> "LampiranChecker":
        return cls(parser, _REQUIRED_LAMPIRAN_VGK, "PKM-VGK")

    @classmethod
    def for_pkm_re(cls, parser: DocxParser) -> "LampiranChecker":
        # PKM-RE: 5 lampiran (TANPA Gambaran Teknologi yang khas KC).
        return cls(parser, _REQUIRED_LAMPIRAN_RE, "PKM-RE")

    @classmethod
    def for_pkm_rsh(cls, parser: DocxParser) -> "LampiranChecker":
        # PKM-RSH: 6 lampiran (DENGAN Kuisioner/Pedoman Wawancara di item #5).
        return cls(parser, _REQUIRED_LAMPIRAN_RSH, "PKM-RSH")

    @classmethod
    def for_pkm_k(cls, parser: DocxParser) -> "LampiranChecker":
        # PKM-K pakai daftar lampiran wajib yang SAMA dengan PKM-VGK (5 lampiran).
        return cls(parser, _REQUIRED_LAMPIRAN_VGK, "PKM-K")

    @classmethod
    def for_pkm_ki(cls, parser: DocxParser) -> "LampiranChecker":
        return cls(parser, _REQUIRED_LAMPIRAN_KI, "PKM-KI")

    @classmethod
    def for_pkm_pi(cls, parser: DocxParser) -> "LampiranChecker":
        return cls(parser, _REQUIRED_LAMPIRAN_PI, "PKM-PI")

    @classmethod
    def for_pkm_pm(cls, parser: DocxParser) -> "LampiranChecker":
        return cls(parser, _REQUIRED_LAMPIRAN_PM, "PKM-PM")

    @classmethod
    def for_pkm_ai(cls, parser: DocxParser) -> "LampiranChecker":
        # PKM-AI tanpa Daftar Lampiran → require_daftar=False.
        return cls(parser, _REQUIRED_LAMPIRAN_AI, "PKM-AI", require_daftar=False)

    @classmethod
    def for_pkm_gft(cls, parser: DocxParser) -> "LampiranChecker":
        # PKM-GFT: Daftar Lampiran gabung di Daftar Isi (tidak ada section
        # terpisah) → require_daftar=False, validasi hanya pada body LAMPIRAN.
        return cls(parser, _REQUIRED_LAMPIRAN_GFT, "PKM-GFT", require_daftar=False)

    # -------------------------------------------------------------------------
    # Public
    # -------------------------------------------------------------------------

    def check(self, index=None, require_funding_letter: bool = False) -> LampiranCheckResult:
        """
        require_funding_letter: bila True, wajibkan lampiran "Surat Pernyataan
        Komitmen Tambahan Pendanaan" di body Lampiran. Orchestrator menyetel ini
        True hanya untuk proposal yang mencantumkan dana Instansi Lain (>Rp0) di
        Rekap Sumber Dana Bab 4. Tidak diteruskan untuk PKM-AI/PKM-GFT (tanpa
        anggaran), sehingga syarat ini otomatis tidak berlaku di sana.
        """
        result = LampiranCheckResult(status="pass")

        # Stage 1: DAFTAR LAMPIRAN ada? — hanya untuk skema yang mewajibkannya.
        # PKM-AI: tidak ada Daftar Lampiran, langsung loncat ke body check.
        if self.require_daftar:
            daftar_start = self._find_daftar_lampiran_idx()
            if daftar_start is None:
                result.status = "fail"
                result.messages.append(CheckMessage(
                    level="fail",
                    text="DAFTAR LAMPIRAN tidak ditemukan",
                ))
                return result
            daftar_corpus = self._collect_daftar_lampiran_text(daftar_start)
        else:
            daftar_corpus = ""  # tidak dipakai

        # Stage 2: korpus body
        body_start = self._find_lampiran_section_start()
        body_text_corpus = self._collect_text_from_lampiran(body_start)

        # OCR body section hanya dipanggil kalau perlu (lazy)
        ocr_corpus_cache: Optional[str] = None

        def body_ocr_text() -> str:
            nonlocal ocr_corpus_cache, index
            if ocr_corpus_cache is not None:
                return ocr_corpus_cache
            if body_start is None:
                ocr_corpus_cache = ""
                return ocr_corpus_cache
            if index is None:
                from app.services.lampiran_index import LampiranOcrIndex
                index = LampiranOcrIndex(self.parser)
            rids = index.rids_in_range(body_start)
            ocr_corpus_cache = index.ocr_text_for_rids(rids) if rids else ""
            return ocr_corpus_cache

        missing_in_body: list[tuple[int, str]] = []
        missing_in_daftar_only: list[tuple[int, str]] = []

        for num, keywords, label in self.required:
            in_body_text = _anchored_keywords_match(keywords, body_text_corpus)
            in_body = in_body_text
            if not in_body:
                ocr_text = body_ocr_text()
                if ocr_text:
                    # OCR fallback: terima anchored ATAU keyword-anywhere
                    # (heading "Lampiran N" di gambar tak selalu ter-OCR jelas).
                    in_body = (
                        _anchored_keywords_match(keywords, ocr_text)
                        or _keywords_in_text(keywords, ocr_text)
                    )

            # Cross-check Daftar Lampiran hanya kalau skema mewajibkannya.
            in_daftar = (
                _anchored_keywords_match(keywords, daftar_corpus)
                if self.require_daftar
                else True
            )

            if not in_body:
                missing_in_body.append((num, label))
            elif not in_daftar:
                missing_in_daftar_only.append((num, label))

        # Stage 2b: Surat Pernyataan Komitmen Tambahan Pendanaan (kondisional) —
        # wajib hanya bila proposal mencantumkan dana Instansi Lain. Cocokkan teks
        # body dulu; baru OCR fallback bila perlu (lazy, reuse cache index).
        funding_letter_missing = False
        if require_funding_letter:
            present = _funding_letter_present(body_text_corpus)
            if not present:
                present = _funding_letter_present(body_ocr_text())
            funding_letter_missing = not present

        # Stage 3: susun pesan
        for num, label in missing_in_body:
            result.messages.append(CheckMessage(
                level="fail",
                text=f'Tidak ditemukan Lampiran {num} "{label}" pada halaman lampiran',
            ))
        for num, label in missing_in_daftar_only:
            result.messages.append(CheckMessage(
                level="fail",
                text=f'Kesalahan kelengkapan Daftar Lampiran, tidak ditemukan "{label}"',
            ))
        if funding_letter_missing:
            result.messages.append(CheckMessage(
                level="fail",
                text=(
                    f'Tidak ditemukan "{_FUNDING_LETTER_LABEL}" pada halaman lampiran '
                    f"(wajib karena proposal mencantumkan dana dari Instansi Lain)"
                ),
            ))

        if result.messages:
            result.status = "fail"
        else:
            result.messages.append(CheckMessage(
                level="pass",
                text=(
                    f"Kelengkapan lampiran {self.schema_label} sesuai: semua "
                    f"{len(self.required)} lampiran wajib ditemukan di Daftar "
                    f"Lampiran dan halaman Lampiran."
                ),
            ))
        return result

    # -------------------------------------------------------------------------
    # Helpers: lokasi section
    # -------------------------------------------------------------------------

    def _find_daftar_lampiran_idx(self) -> Optional[int]:
        for p in self.parser.paragraphs:
            if _DAFTAR_LAMPIRAN_RE.match(p.text.strip()):
                return p.index
        return None

    def _find_lampiran_section_start(self) -> Optional[int]:
        for p in self.parser.paragraphs:
            t = p.text.strip()
            if _LAMPIRAN_SECTION_RE.match(t) and p.is_heading:
                return p.index
        return None

    # -------------------------------------------------------------------------
    # Helpers: korpus teks
    # -------------------------------------------------------------------------

    def _collect_daftar_lampiran_text(self, daftar_idx: int) -> str:
        """
        Teks entri di blok Daftar Lampiran — sampai heading section lain.

        Pendekatan ganda:
        1. Coba dari parser.paragraphs (cepat, lazim).
        2. Kalau hasil kosong/sangat pendek (kasus TOC field code yang berada di
           dalam <w:sdt> dan TIDAK dienumerasi sebagai paragraf python-docx),
           jatuh ke ekstraksi XML body langsung — termasuk konten SDT.
        """
        parts: list[str] = []
        next_heading_idx: Optional[int] = None
        for p in self.parser.paragraphs:
            if p.index <= daftar_idx:
                continue
            t = p.text.strip()
            if p.is_heading and t and not _LAMPIRAN_N_RE.search(t):
                next_heading_idx = p.index
                break
            if t:
                parts.append(p.text)
        result = " ".join(parts)

        # Fallback XML kalau kosong (TOC field di <w:sdt>) — extract dari body XML.
        if not result.strip():
            result = self._collect_daftar_lampiran_text_xml(daftar_idx, next_heading_idx)
        return result

    def _collect_daftar_lampiran_text_xml(
        self, daftar_idx: int, next_heading_idx: Optional[int]
    ) -> str:
        """Extract teks dari semua node setelah paragraf DAFTAR LAMPIRAN sampai
        paragraf heading berikutnya (eksklusif). Iterasi anak-langsung <w:body>,
        ambil semua <w:t> dalam <w:p> dan <w:sdt> (TOC field)."""
        from docx.oxml.ns import qn
        try:
            body = self.parser.doc.element.body
        except Exception:
            return ""
        para_counter = -1
        parts: list[str] = []

        def _extract_text(el) -> str:
            # Pisah dengan spasi: hindari "Pendamping11Lampiran" yang menghilangkan
            # word-boundary di regex _LAMPIRAN_N_RE.
            return " ".join(
                (t.text or "").strip()
                for t in el.iter(qn("w:t"))
                if (t.text or "").strip()
            )

        for child in body:
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag == "p":
                para_counter += 1
                if para_counter <= daftar_idx:
                    continue
                if next_heading_idx is not None and para_counter >= next_heading_idx:
                    break
                text = _extract_text(child)
                if text:
                    parts.append(text)
            elif tag == "sdt":
                # SDT (TOC field) berada di antara paragraf — counter paragraf tidak
                # bergerak. Hanya ambil setelah daftar_idx sudah lewat.
                if para_counter < daftar_idx:
                    continue
                if next_heading_idx is not None and para_counter >= next_heading_idx:
                    break
                text = _extract_text(child)
                if text:
                    parts.append(text)
        return " ".join(parts)

    def _collect_text_from_lampiran(self, start_idx: Optional[int]) -> str:
        """Kumpulkan teks paragraf mulai dari heading 'LAMPIRAN' body s.d. akhir."""
        if start_idx is None:
            return ""
        parts: list[str] = []
        for p in self.parser.paragraphs:
            if p.index < start_idx:
                continue
            if p.text.strip():
                parts.append(p.text)
        return " ".join(parts)


# =============================================================================
# Helpers: matching kata kunci
# =============================================================================


def _anchored_keywords_match(keywords: list[str], corpus: str) -> bool:
    """
    Tiap "Lampiran N" di corpus → window 200 char setelahnya; match kalau
    SEMUA keyword muncul di window. Tidak tergantung nomor N — agar variasi
    penomoran lapangan tidak menggagalkan match.
    """
    if not corpus or not keywords:
        return False
    cl = corpus.lower()
    for m in _LAMPIRAN_N_RE.finditer(cl):
        window = cl[m.start(): m.start() + 200]
        if all(kw in window for kw in keywords):
            return True
    return False


def _funding_letter_present(corpus: str) -> bool:
    """
    True bila korpus memuat salah satu frasa penanda surat komitmen tambahan
    pendanaan (lihat _FUNDING_LETTER_PHRASES). Frasa sengaja semuanya mengandung
    "pendanaan" agar penyebutan "Instansi Lain" biasa di biodata/RAB tidak
    memicu false positive.
    """
    if not corpus:
        return False
    cl = corpus.lower()
    return any(p in cl for p in _FUNDING_LETTER_PHRASES)


def _keywords_in_text(keywords: list[str], corpus: str) -> bool:
    """
    Fallback longgar untuk OCR: cukup semua keyword muncul di corpus, tanpa
    anchor "Lampiran N" (kadang heading di gambar tak ter-OCR jelas, tapi isi
    form-nya jelas terbaca dan memuat kata kunci judul lampiran).
    """
    if not corpus or not keywords:
        return False
    cl = corpus.lower()
    return all(kw in cl for kw in keywords)
