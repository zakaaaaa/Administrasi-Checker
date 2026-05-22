"""
Test suite untuk ScheduleChecker (tabel jadwal kegiatan).

Cara jalankan:
    python3 -m unittest tests.test_schedule_checker -v
"""

import os
import tempfile
import unittest
from pathlib import Path

from docx import Document

from app.services.docx_parser import DocxParser
from app.services.schedule_checker import (
    ScheduleChecker,
    _dedup_consecutive,
    _norm,
)
from app.services.schedule_rules import get_pkm_kc_schedule_rules

REF_FILE = Path("/Users/zakakurniarahman/Documents/AdministrasiChecker/format tabel.docx")


def _build_schedule_docx(headers, months, pic_rows) -> str:
    """
    Buat docx berisi satu tabel jadwal.

    headers: [label_no, label_kegiatan, label_bulan, label_pic]
    months : list angka bulan untuk sub-header
    pic_rows: list PIC per kegiatan ('' = kosong)
    """
    doc = Document()
    doc.add_paragraph("Tabel 3. Jadwal Kegiatan PKM")
    ncols = 2 + len(months) + 1
    t = doc.add_table(rows=2 + len(pic_rows), cols=ncols)
    t.cell(0, 0).text = headers[0]
    t.cell(0, 1).text = headers[1]
    for c in range(len(months)):
        t.cell(0, 2 + c).text = headers[2]
    t.cell(0, ncols - 1).text = headers[3]
    for c, mn in enumerate(months):
        t.cell(1, 2 + c).text = str(mn)
    for i, pic in enumerate(pic_rows):
        r = 2 + i
        t.cell(r, 0).text = str(i + 1)
        t.cell(r, 1).text = f"Kegiatan {i + 1}"
        t.cell(r, ncols - 1).text = pic
    f = tempfile.mktemp(suffix=".docx")
    doc.save(f)
    return f

_GOOD = ["No", "Jadwal Kegiatan", "Bulan", "Penanggung Jawab"]


def _run(headers, months, pic_rows):
    f = _build_schedule_docx(headers, months, pic_rows)
    try:
        return ScheduleChecker.for_pkm_kc(DocxParser(f)).check()
    finally:
        os.remove(f)


# ============================================================================
# Helpers
# ============================================================================


class TestHelpers(unittest.TestCase):
    def test_norm_collapses_whitespace_and_case(self):
        self.assertEqual(_norm("  Penanggung   Jawab "), "penanggung jawab")

    def test_dedup_drops_empty_and_merge_duplicates(self):
        row = ["No", "Bulan", "Bulan", "Bulan", "", "Penanggung Jawab"]
        self.assertEqual(_dedup_consecutive(row), ["No", "Bulan", "Penanggung Jawab"])


# ============================================================================
# Rules
# ============================================================================


class TestRules(unittest.TestCase):
    def test_pkm_kc_defaults(self):
        r = get_pkm_kc_schedule_rules()
        self.assertEqual(r.expected_headers, _GOOD)
        self.assertEqual(r.required_months, 4)
        self.assertTrue(r.require_pic_filled)


# ============================================================================
# Header strict
# ============================================================================


class TestHeaders(unittest.TestCase):
    def test_all_correct_passes(self):
        res = _run(_GOOD, [1, 2, 3, 4], ["A", "B"])
        self.assertEqual(res.status, "pass")

    def test_case_insensitive_header_passes(self):
        res = _run(["No", "JADWAL KEGIATAN", "Bulan", "PENANGGUNG JAWAB"], [1, 2, 3, 4], ["A"])
        self.assertEqual(res.status, "pass")

    def test_pic_label_rejected(self):
        res = _run(["No", "Jadwal Kegiatan", "Bulan", "PIC"], [1, 2, 3, 4], ["A"])
        self.assertEqual(res.status, "fail")
        blob = " ".join(m.text for m in res.messages)
        self.assertIn("PIC", blob)
        self.assertIn("Penanggung Jawab", blob)

    def test_pj_abbreviation_rejected(self):
        res = _run(["No", "Jadwal Kegiatan", "Bulan", "PJ"], [1, 2, 3, 4], ["A"])
        self.assertEqual(res.status, "fail")

    def test_kegiatan_synonym_rejected(self):
        res = _run(["No", "Jenis Kegiatan", "Bulan", "Penanggung Jawab"], [1, 2, 3, 4], ["A"])
        self.assertEqual(res.status, "fail")
        blob = " ".join(m.text for m in res.messages)
        self.assertIn("Jenis Kegiatan", blob)


# ============================================================================
# Rentang bulan
# ============================================================================


class TestMonths(unittest.TestCase):
    def test_exactly_four_months_ok(self):
        res = _run(_GOOD, [1, 2, 3, 4], ["A"])
        self.assertEqual(res.status, "pass")

    def test_five_months_fail(self):
        res = _run(_GOOD, [1, 2, 3, 4, 5], ["A"])
        self.assertEqual(res.status, "fail")
        self.assertTrue(any("bulan" in m.text.lower() and m.level == "fail" for m in res.messages))

    def test_three_months_fail(self):
        res = _run(_GOOD, [1, 2, 3], ["A"])
        self.assertEqual(res.status, "fail")
        self.assertTrue(any("tepat 4 bulan" in m.text.lower() for m in res.messages))

    def test_two_months_fail(self):
        """Regresi: file copy NUKI dikurangi jadi 2 bulan → harus fail."""
        res = _run(_GOOD, [1, 2], ["A"])
        self.assertEqual(res.status, "fail")
        self.assertTrue(any("jumlah bulan" in m.text.lower() for m in res.messages))


# ============================================================================
# Penanggung jawab terisi
# ============================================================================


class TestPicFilled(unittest.TestCase):
    def test_empty_pic_is_warning(self):
        res = _run(_GOOD, [1, 2, 3, 4], ["A", "", "C"])
        self.assertEqual(res.status, "warning")
        self.assertTrue(any("nomor 2" in m.text for m in res.messages))

    def test_all_filled_passes(self):
        res = _run(_GOOD, [1, 2, 3, 4], ["A", "B", "C"])
        self.assertEqual(res.status, "pass")


# ============================================================================
# Tabel tidak ditemukan
# ============================================================================


class TestTableNotFound(unittest.TestCase):
    def test_no_table_is_warning(self):
        doc = Document()
        doc.add_paragraph("Dokumen tanpa tabel jadwal.")
        f = tempfile.mktemp(suffix=".docx")
        doc.save(f)
        try:
            res = ScheduleChecker.for_pkm_kc(DocxParser(f)).check()
        finally:
            os.remove(f)
        self.assertEqual(res.status, "warning")
        self.assertTrue(any("tidak terdeteksi" in m.text.lower() for m in res.messages))


# ============================================================================
# Dokumen acuan asli
# ============================================================================


class TestReferenceDoc(unittest.TestCase):
    def test_reference_template_passes(self):
        if not REF_FILE.exists():
            self.skipTest(f"{REF_FILE.name} tidak ada")
        res = ScheduleChecker.for_pkm_kc(DocxParser(REF_FILE)).check()
        self.assertEqual(res.status, "pass", [m.text for m in res.messages])

    def test_to_dict_serializable(self):
        res = _run(_GOOD, [1, 2, 3, 4], ["A"])
        d = res.to_dict()
        self.assertIn("status", d)
        self.assertIn("messages", d)
        self.assertIsInstance(d["messages"], list)


if __name__ == "__main__":
    unittest.main(verbosity=2)
