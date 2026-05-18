"""
Test suite untuk FormatChecker.

Cara jalankan:
    python3 -m unittest tests.test_format_checker -v
"""

import unittest
import tempfile
from pathlib import Path
from unittest.mock import MagicMock, patch

from docx import Document
from lxml import etree

from app.services.docx_parser import DocxParser, ParagraphInfo
from app.services.format_checker import (
    FormatChecker,
    FormatRules,
    FOREIGN_WORDS,
    _is_figure_table_caption_paragraph,
    get_pkm_format_rules,
)
from app.services.style_resolver import StyleResolver, ResolvedFont
from app.services.schema_rules import get_pkm_kc_proposal_rules

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

SAMPLE_DIR = Path(__file__).parent / "sample_docs"
DUMMY_FILE = SAMPLE_DIR / "dummy_pkm_kc.docx"
REAL_FILE = SAMPLE_DIR / "A410170082.docx"


# ============================================================================
# Test: FormatRules construction
# ============================================================================


class TestFormatRules(unittest.TestCase):
    def test_default_rules(self):
        r = get_pkm_format_rules()
        self.assertEqual(r.font_name, "Times New Roman")
        self.assertEqual(r.font_size_pt, 12.0)
        self.assertEqual(r.margin_left_cm, 4.0)
        self.assertEqual(r.margin_right_cm, 3.0)
        self.assertEqual(r.paper_width_cm, 21.0)
        self.assertEqual(r.paper_height_cm, 29.7)
        self.assertEqual(r.line_spacing, 1.15)


# ============================================================================
# Test: StyleResolver
# ============================================================================


class TestStyleResolverDummy(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        if not DUMMY_FILE.exists():
            raise unittest.SkipTest("Dummy belum di-generate")
        cls.parser = DocxParser(DUMMY_FILE)
        cls.resolver = StyleResolver(cls.parser)

    def test_resolves_default_font(self):
        """Default style 'Normal' di dummy = Times New Roman 12pt."""
        # Cari paragraf body (bukan heading)
        body_paras = [p for p in self.parser.paragraphs
                      if not p.is_heading and p.text.strip()]
        self.assertGreater(len(body_paras), 0)
        font = self.resolver.resolve_paragraph_font(body_paras[0].index)
        # Font name harus ke-resolve sebagai TNR (set di Normal style oleh build_dummy)
        self.assertEqual(font.name, "Times New Roman",
                         f"Got {font.name} (source={font.name_source})")

    def test_returns_resolved_font_object(self):
        font = self.resolver.resolve_paragraph_font(0)
        self.assertIsInstance(font, ResolvedFont)

    def test_resolves_office_theme_font_name(self):
        """Theme font di dokumen harus terbaca sebagai nama font, bukan None."""
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "theme_calibri.docx"
            doc = Document()
            doc.add_paragraph("Mendukung deteksi font theme.")
            doc.save(path)

            parser = DocxParser(path)
            resolver = StyleResolver(parser)
            font = resolver.resolve_paragraph_font(0)

        self.assertIsNotNone(font.name)
        self.assertIn("theme", font.name_source)

    def test_theme_fallback_maps_minor_hansi_to_calibri(self):
        """Jika theme file tidak tersedia, minorHAnsi tetap dianggap Calibri."""
        parser = MagicMock()
        parser.read_raw_part.return_value = None
        resolver = StyleResolver(parser)
        self.assertEqual(resolver._resolve_theme_font("minorHAnsi"), "Calibri")

    def test_east_asia_theme_does_not_override_latin_font(self):
        """eastAsiaTheme tidak boleh membuat teks Latin TNR terbaca Calibri."""
        parser = MagicMock()
        parser.read_raw_part.return_value = None
        resolver = StyleResolver(parser)
        font = ResolvedFont()
        rpr = etree.fromstring(
            f'<w:rPr xmlns:w="{W_NS}"><w:rFonts w:eastAsiaTheme="minorHAnsi"/></w:rPr>'
        )
        resolver._fill_from_rpr(font, rpr, source="paragraph")
        self.assertIsNone(font.name)


class TestStyleResolverReal(unittest.TestCase):
    """
    Test resolver pada dokumen real — kasus yang motivasi bikin resolver:
    772 run "font tidak terbaca" di parser. Resolver harus bisa resolve
    semuanya ke font Normal style.
    """

    @classmethod
    def setUpClass(cls):
        if not REAL_FILE.exists():
            raise unittest.SkipTest(f"{REAL_FILE.name} tidak ada")
        cls.parser = DocxParser(REAL_FILE)
        cls.resolver = StyleResolver(cls.parser)

    def test_resolver_returns_font_for_body_paras(self):
        """Untuk paragraf body, resolver harus return font yang valid (bukan None)."""
        sample_indices = [170, 175, 213, 260, 300]
        resolved_count = 0
        for idx in sample_indices:
            if idx >= len(self.parser.paragraphs):
                continue
            font = self.resolver.resolve_paragraph_font(idx)
            if font.name is not None:
                resolved_count += 1
        self.assertGreater(resolved_count, 0,
            "Resolver tidak bisa resolve font untuk satupun body paragraph")


# ============================================================================
# Test: FormatChecker pada DUMMY (compliant)
# ============================================================================


class TestFormatCheckerOnDummy(unittest.TestCase):
    """
    Dummy dirancang compliant: TNR 12pt, margin L=4 R/T/B=3, line spacing 1.15.
    Tapi paper size dummy adalah Letter (21.59x27.94), BUKAN A4 → harus FAIL
    di paper_size check (insight ini muncul di iterasi DocxParser).
    """

    @classmethod
    def setUpClass(cls):
        if not DUMMY_FILE.exists():
            raise unittest.SkipTest("Dummy belum di-generate")
        cls.parser = DocxParser(DUMMY_FILE)
        cls.result = FormatChecker(cls.parser).check()

    def test_returns_result(self):
        self.assertIsNotNone(self.result)
        self.assertIn("paper_size", self.result.checks)
        self.assertIn("margin", self.result.checks)
        self.assertIn("font_body", self.result.checks)

    def test_paper_size_fails_letter(self):
        """Dummy pakai Letter (21.59×27.94), harus FAIL untuk A4."""
        self.assertEqual(self.result.checks["paper_size"].status, "fail")

    def test_margin_passes(self):
        """Dummy margin: L=4, R/T/B=3 → match aturan PKM."""
        self.assertEqual(self.result.checks["margin"].status, "pass")

    def test_font_body_passes(self):
        """Dummy default Normal style = TNR 12pt → font_body harus pass."""
        sec = self.result.checks["font_body"]
        if sec.status != "pass":
            print(f"\nDEBUG font issues: {[i.issue for i in sec.issues[:5]]}")
        self.assertEqual(sec.status, "pass")

    def test_to_dict_serializable(self):
        d = self.result.to_dict()
        self.assertIn("status", d)
        self.assertIn("checks", d)
        self.assertIn("messages", d)


# ============================================================================
# Test: FormatChecker pada DOKUMEN REAL
# ============================================================================


class TestFormatCheckerOnRealDoc(unittest.TestCase):
    """
    Dokumen real `A410170082.docx`:
    - Margin L=0, R=0 di semua section → harus FAIL margin
    - Paper W=21.008 H=29.704 → seharusnya pass A4 (dalam toleransi 0.1cm)
    """

    @classmethod
    def setUpClass(cls):
        if not REAL_FILE.exists():
            raise unittest.SkipTest(f"{REAL_FILE.name} tidak ada")
        cls.parser = DocxParser(REAL_FILE)
        cls.result = FormatChecker(cls.parser).check()

    def test_overall_status_is_fail(self):
        """Margin 0 sudah cukup membuat overall fail."""
        self.assertEqual(self.result.status, "fail")

    def test_margin_fails(self):
        sec = self.result.checks["margin"]
        self.assertEqual(sec.status, "fail")
        # Pastikan ada flag untuk margin left/right
        issues_text = " ".join(i.issue for i in sec.issues)
        self.assertIn("left", issues_text)

    def test_paper_size_passes_a4(self):
        """A410170082 paper 21.008x29.704 → dalam toleransi A4."""
        sec = self.result.checks["paper_size"]
        self.assertEqual(sec.status, "pass",
            f"Issues: {[i.issue for i in sec.issues[:3]]}")

    def test_font_body_passes_or_minimal_issues(self):
        """
        Body teks dokumen real pakai TNR 12pt (verifikasi manual di
        StyleResolver). Beberapa paragraf bisa pakai size berbeda
        (header-like content), tapi mayoritas TNR 12pt.
        """
        sec = self.result.checks["font_body"]
        # Ada beberapa size non-12pt yang sudah kita lihat (20pt, 13pt, 10pt)
        # tapi mayoritas harus 12pt
        size_dist = sec.detail.get("size_distribution", {})
        # 12.0pt harus jadi distribusi terbanyak
        if size_dist:
            most_common_size = max(size_dist.items(), key=lambda x: x[1])[0]
            self.assertEqual(most_common_size, 12.0,
                f"Distribusi size: {size_dist}")


# ============================================================================
# Test: deteksi foreign words italic
# ============================================================================


class TestForeignWordsDetection(unittest.TestCase):
    """Test deteksi kata asing yang tidak italic, pakai dummy."""

    @classmethod
    def setUpClass(cls):
        if not DUMMY_FILE.exists():
            raise unittest.SkipTest("Dummy belum di-generate")
        cls.parser = DocxParser(DUMMY_FILE)
        cls.result = FormatChecker(cls.parser).check()

    def test_iot_in_dummy_flagged(self):
        """
        Dummy sengaja punya 'Internet of Things' tanpa italic di paragraf 12.
        Harus terdeteksi di foreign_words_italic check.
        """
        sec = self.result.checks["foreign_words_italic"]
        # Minimal 1 issue terdeteksi
        self.assertGreater(len(sec.issues), 0,
            "Foreign words checker tidak detect 'Internet of Things' di dummy")

    def test_foreign_words_dictionary_not_empty(self):
        self.assertGreater(len(FOREIGN_WORDS), 10)
        self.assertIn("internet of things", FOREIGN_WORDS)


class TestCaptionParagraphDetection(unittest.TestCase):
    """Caption gambar/tabel dikecualikan dari wajib justify."""

    def test_figure_caption_indonesian(self):
        self.assertTrue(
            _is_figure_table_caption_paragraph(
                "Gambar 2. Logo Protextify untuk sistem keamanan data."
            )
        )

    def test_table_caption(self):
        self.assertTrue(
            _is_figure_table_caption_paragraph("Tabel 1 — Ringkasan hasil uji")
        )

    def test_not_caption_body_text(self):
        self.assertFalse(
            _is_figure_table_caption_paragraph(
                "Pada gambar 2 terlihat bahwa hasil penelitian menunjukkan "
                "tren positif yang berkelanjutan."
            )
        )

    def test_long_line_not_treated_as_caption(self):
        long_tail = "x" * 300
        self.assertFalse(
            _is_figure_table_caption_paragraph("Gambar 2. " + long_tail)
        )


class TestAlignmentSkipsFigureCaptions(unittest.TestCase):
    """Alignment: center pada caption panjang tidak boleh false positive."""

    def test_center_caption_no_alignment_issue(self):
        from app.services import format_checker as fc

        paras = [
            ParagraphInfo(
                index=0,
                text="Gambar 2. Logo Protextify untuk sistem keamanan data.",
                is_heading=False,
                alignment="center",
            ),
        ]
        parser = MagicMock()
        parser.paragraphs = paras
        with patch.object(fc, "StyleResolver", return_value=MagicMock()):
            checker = FormatChecker(parser)
            sec = checker._check_alignment()
        self.assertEqual(sec.status, "pass")
        self.assertEqual(len(sec.issues), 0)

    def test_center_body_paragraph_still_fails(self):
        from app.services import format_checker as fc

        paras = [
            ParagraphInfo(
                index=0,
                text="Ini adalah paragraf body biasa yang cukup panjang.",
                is_heading=False,
                alignment="center",
            ),
        ]
        parser = MagicMock()
        parser.paragraphs = paras
        with patch.object(fc, "StyleResolver", return_value=MagicMock()):
            checker = FormatChecker(parser)
            sec = checker._check_alignment()
        self.assertEqual(sec.status, "fail")
        self.assertGreaterEqual(len(sec.issues), 1)


# ============================================================================
# Test: synthetic FormatRules behavior
# ============================================================================


class TestRulesEdgeCases(unittest.TestCase):
    def test_disable_justify_check(self):
        """Kalau require_justify=False, sub-check alignment tidak di-run."""
        if not DUMMY_FILE.exists():
            self.skipTest("Dummy belum di-generate")
        parser = DocxParser(DUMMY_FILE)
        rules = FormatRules(require_justify=False)
        result = FormatChecker(parser, rules).check()
        self.assertNotIn("alignment", result.checks)


if __name__ == "__main__":
    unittest.main(verbosity=2)
