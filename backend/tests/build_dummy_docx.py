"""
Generate dummy .docx PKM-KC sederhana untuk smoke testing DocxParser.

Dokumen yang dihasilkan punya:
- Beberapa heading section (DAFTAR ISI, BAB 1, BAB 2, ..., DAFTAR PUSTAKA, LAMPIRAN)
- Body teks dengan font Times New Roman 12pt
- Satu tabel sederhana (mock RAB)
- Section break (untuk test multi-section parsing)
- Header & footer (untuk test header_xmls/footer_xmls)
"""

from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION


def build_dummy_pkm_kc(out_path: Path) -> None:
    doc = Document()

    # Set margin section default sesuai PKM (kiri 4cm, lainnya 3cm)
    section = doc.sections[0]
    section.left_margin = Cm(4)
    section.right_margin = Cm(3)
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)

    # Set default style font Times New Roman 12pt
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # ----- Front matter (zona awal — harusnya romawi pojok kanan bawah) -----
    h = doc.add_heading("DAFTAR ISI", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("BAB 1. PENDAHULUAN ............... 1")
    doc.add_paragraph("BAB 2. TINJAUAN PUSTAKA .......... 3")
    doc.add_paragraph("BAB 3. TAHAP PELAKSANAAN ......... 5")
    doc.add_paragraph("BAB 4. BIAYA DAN JADWAL .......... 7")
    doc.add_paragraph("DAFTAR PUSTAKA ................... 9")

    doc.add_heading("DAFTAR LAMPIRAN", level=1)
    doc.add_paragraph("Lampiran 1. Biodata Ketua dan Anggota")
    doc.add_paragraph("Lampiran 2. Justifikasi Anggaran")

    # Section break untuk pindah ke bagian inti (zona arab pojok kanan atas)
    doc.add_section(WD_SECTION.NEW_PAGE)

    # ----- Bagian inti -----
    doc.add_heading("BAB 1. PENDAHULUAN", level=1)
    p = doc.add_paragraph(
        "Latar belakang penelitian ini dimulai dari permasalahan "
        "energi terbarukan. Menurut Smith (2020) energi surya menjadi "
        "alternatif yang paling menjanjikan."
    )
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.15

    # Sengaja tambahkan satu run dengan kata asing yang TIDAK italic
    # supaya FormatChecker nanti bisa men-deteksinya
    p2 = doc.add_paragraph()
    r1 = p2.add_run("Penggunaan teknologi ")
    r1.font.name = "Times New Roman"
    r1.font.size = Pt(12)
    r2 = p2.add_run("Internet of Things")  # asing, sengaja TIDAK italic
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(12)
    r3 = p2.add_run(" sangat relevan untuk smart farming.")
    r3.font.name = "Times New Roman"
    r3.font.size = Pt(12)

    doc.add_heading("BAB 2. TINJAUAN PUSTAKA", level=1)
    doc.add_paragraph("Tinjauan pustaka membahas teori-teori pendukung.")

    doc.add_heading("BAB 3. TAHAP PELAKSANAAN", level=1)
    doc.add_paragraph("Tahap pelaksanaan terdiri dari empat fase.")

    doc.add_heading("BAB 4. BIAYA DAN JADWAL KEGIATAN", level=1)
    doc.add_paragraph("Berikut adalah ringkasan biaya:")

    # Tabel RAB sederhana — Bab 4 (4 kategori standar)
    tbl = doc.add_table(rows=5, cols=3)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    hdr[0].text = "No"
    hdr[1].text = "Jenis Pengeluaran"
    hdr[2].text = "Total (Rp)"

    rows_data = [
        ("1", "Bahan habis pakai", "4.500.000"),
        ("2", "Sewa dan jasa", "1.200.000"),
        ("3", "Transportasi lokal", "1.000.000"),
        ("4", "Lain-lain", "300.000"),
    ]
    for i, (no, jenis, total) in enumerate(rows_data, start=1):
        tbl.rows[i].cells[0].text = no
        tbl.rows[i].cells[1].text = jenis
        tbl.rows[i].cells[2].text = total

    doc.add_heading("DAFTAR PUSTAKA", level=1)
    # Sengaja sertakan satu entry dengan "et al." supaya ReferenceValidator nanti
    # bisa men-deteksinya sebagai pelanggaran Harvard strict
    doc.add_paragraph(
        "Smith, J. (2020) Solar Energy Fundamentals. New York: Springer."
    )
    doc.add_paragraph(
        "Rahmadi et al. (2022) 'Smart farming review', "
        "Jurnal Teknologi, 15(2), pp. 100-115."
    )

    doc.add_heading("LAMPIRAN", level=1)
    doc.add_heading("Lampiran 1. Biodata Ketua", level=2)
    doc.add_paragraph("Nama Lengkap: Muhammad Ibnu Pratama")
    doc.add_paragraph("NIM: 2114321023")
    doc.add_paragraph("Tahun: 2027")

    doc.add_heading("Lampiran 2. Justifikasi Anggaran", level=2)
    # Tabel rinci Lampiran 2 — sengaja mismatch dengan Bab 4 untuk test cross-check
    tbl2 = doc.add_table(rows=4, cols=4)
    tbl2.style = "Table Grid"
    hdr2 = tbl2.rows[0].cells
    hdr2[0].text = "No"
    hdr2[1].text = "Item"
    hdr2[2].text = "Kategori"
    hdr2[3].text = "Total (Rp)"
    items = [
        ("1", "Sensor DHT22", "Bahan habis pakai", "4.750.000"),  # mismatch sengaja
        ("2", "Sewa lab", "Sewa dan jasa", "1.200.000"),
        ("3", "Transport survey", "Transportasi lokal", "1.000.000"),
    ]
    for i, (no, item, kategori, total) in enumerate(items, start=1):
        tbl2.rows[i].cells[0].text = no
        tbl2.rows[i].cells[1].text = item
        tbl2.rows[i].cells[2].text = kategori
        tbl2.rows[i].cells[3].text = total

    doc.save(out_path)


if __name__ == "__main__":
    out = Path(__file__).parent / "sample_docs" / "dummy_pkm_kc.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    build_dummy_pkm_kc(out)
    print(f"Saved: {out}  ({out.stat().st_size} bytes)")